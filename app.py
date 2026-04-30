from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, Response
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import os
import urllib.parse
import uuid
from datetime import datetime, timedelta
import pytz
import sys
import csv
import io
import requests
from PIL import Image
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import base64
import json
import re
from collections import defaultdict

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XLImage
    OPENPYXL_AVAILABLE = True
except Exception:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    PYDOCX_AVAILABLE = True
except Exception:
    PYDOCX_AVAILABLE = False


app = Flask(__name__)
app.secret_key = 'sk_ons_warehouse_secret_key_2025'
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

# 세션 설정 강화
app.permanent_session_lifetime = timedelta(hours=8)
app.config['SESSION_COOKIE_SECURE'] = False
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_NAME'] = 'sk_ons_session'

# 업로드 폴더 생성
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# 환경변수 확인
DATABASE_URL = os.environ.get('SUPABASE_DB_URL')
ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'Onsn1103813!')
SUPABASE_URL = os.environ.get('SUPABASE_URL')
SUPABASE_SERVICE_KEY = os.environ.get('SUPABASE_SERVICE_KEY')

# 이메일 설정 (환경변수에서 가져오기)
SMTP_SERVER = os.environ.get('SMTP_SERVER', 'smtp.gmail.com')
SMTP_PORT = int(os.environ.get('SMTP_PORT', 587))
SMTP_USERNAME = os.environ.get('SMTP_USERNAME')
SMTP_PASSWORD = os.environ.get('SMTP_PASSWORD')

print("=" * 60)
print("🚀 SK오앤에스 DIY System 시작")
print("=" * 60)

# Supabase 연결 필수 체크
if not DATABASE_URL or not DATABASE_URL.startswith('postgresql://'):
    print("❌ 치명적 오류: 올바른 SUPABASE_DB_URL 환경변수가 설정되지 않았습니다!")
    print("📋 해결 방법:")
    print("   1. Render 대시보드에서 Environment Variables 설정")
    print("   2. SUPABASE_DB_URL 추가 (postgresql://로 시작해야 함)")
    print("   3. 재배포")
    print(f"   현재값: {DATABASE_URL[:30] if DATABASE_URL else 'None'}...")
    print("=" * 60)
    sys.exit(1)

print(f"✅ SUPABASE_DB_URL: {DATABASE_URL[:50]}...")
print(f"✅ SUPABASE_URL: {SUPABASE_URL}")

# 허용된 파일 확장자
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}

# 공개 식별자와 DB 저장값 분리
DIY_ACTIVE_SLUG = 'cooling-maintenance'
DIY_ACTIVE_LABEL = '냉방기 예방점검'
DIY_FACILITY_SLUG = 'facility-suitability'
DIY_FACILITY_LABEL = '시설물 적합조사'
DIY_PREPARING_LABEL = '준비중'
DB_ACTIVE_WAREHOUSE = '보라매창고'
DB_FACILITY_WAREHOUSE = '시설물적합조사'
DIY_CHECKLIST_CATEGORY = 'DIY점검'
INSPECTION_METHOD_BUCKET = 'warehouse-photos'
INSPECTION_METHOD_PREFIX = 'inspection-methods'
WAREHOUSES = [DIY_ACTIVE_SLUG, DIY_FACILITY_SLUG]
DEFAULT_EQUIPMENT_PREFIX = 'A05'
DEFAULT_EQUIPMENT_START = 1
EQUIPMENT_NO_PATTERN = re.compile(r'^(?P<prefix>[A-Za-z0-9]+)-(?P<index>\d+)$')


def get_db_warehouse_from_slug(warehouse_slug):
    if warehouse_slug == DIY_ACTIVE_SLUG:
        return DB_ACTIVE_WAREHOUSE
    if warehouse_slug == DIY_FACILITY_SLUG:
        return DB_FACILITY_WAREHOUSE
    return None


def get_display_warehouse_from_slug(warehouse_slug):
    if warehouse_slug == DIY_ACTIVE_SLUG:
        return DIY_ACTIVE_LABEL
    if warehouse_slug == DIY_FACILITY_SLUG:
        return DIY_FACILITY_LABEL
    return DIY_PREPARING_LABEL


def get_slug_from_db_warehouse(warehouse_name):
    if warehouse_name == DB_ACTIVE_WAREHOUSE:
        return DIY_ACTIVE_SLUG
    if warehouse_name == DB_FACILITY_WAREHOUSE:
        return DIY_FACILITY_SLUG
    return DIY_ACTIVE_SLUG


def normalize_warehouse_filter(warehouse_value):
    if not warehouse_value:
        return None
    if warehouse_value in [DIY_ACTIVE_SLUG, DIY_ACTIVE_LABEL, DB_ACTIVE_WAREHOUSE]:
        return DB_ACTIVE_WAREHOUSE
    if warehouse_value in [DIY_FACILITY_SLUG, DIY_FACILITY_LABEL, DB_FACILITY_WAREHOUSE]:
        return DB_FACILITY_WAREHOUSE
    return None


def get_diy_intro_text(warehouse_slug):
    if warehouse_slug == DIY_FACILITY_SLUG:
        return 'DIY작업관리 시설물 적합조사를 진행합니다.'
    return '냉방기 예방점검 대상 점검을 진행합니다'


def uses_equipment_units(warehouse_slug):
    return warehouse_slug == DIY_ACTIVE_SLUG


FACILITY_SURVEY_VERSION = 1
FACILITY_RESULT_OPTIONS = {'fit', 'unfit'}
FACILITY_CHECK_OPTIONS = {'fit', 'unfit', 'na'}


def is_facility_suitability(warehouse_slug):
    return warehouse_slug == DIY_FACILITY_SLUG


FACILITY_SURVEY_SECTIONS = [
    {
        'code': '1',
        'title': '보호기 및 접지(규정 제7조)',
        'checks': [
            '보호기 설치 여부',
            '보호기의 기술기준 준수 여부',
            '접지저항 적합 여부',
        ],
        'detail_fields': [
            {'key': 'ground_measure_count', 'label': '연 ( )회 측정'},
            {'key': 'ground_measure_date', 'label': '( . . . 시행)'},
            {'key': 'ground_ohm', 'label': '( Ω )'},
            {'key': 'opening_date', 'label': '개국일'},
            {'key': 'approval_date', 'label': '사용승인일 ( . . . )'},
            {'key': 'floor_area_m2', 'label': '연면적 ( ㎡ )'},
        ],
    },
    {
        'code': '2',
        'title': '예비전원설비(규정 제10조)',
        'checks': [
            '축전지 및 발전기 설치 여부(사업용)',
            '예비전원 설치 여부(사업용 외 국선 10회선 이상)',
        ],
        'detail_fields': [
            {'key': 'generator_kw', 'label': '비상발전기 용량 (KW)'},
            {'key': 'fuel_tank_total_l', 'label': '유류탱크 용량 총량 (L)'},
            {'key': 'fuel_tank_have_l', 'label': '유류탱크 보유량 (L)'},
            {'key': 'backup_hours', 'label': '백업 가능 시간'},
            {'key': 'rectifier_count', 'label': '정류기 (대)'},
            {'key': 'rectifier_battery_ah', 'label': '정류기 축전지 (AH)'},
            {'key': 'rectifier_max_load_a', 'label': '정류기 최대부하량 (A)'},
            {'key': 'rectifier_avg_hr', 'label': '정류기 평균 (HR)'},
            {'key': 'ups_count', 'label': 'UPS (대)'},
            {'key': 'ups_battery_ah', 'label': 'UPS 축전지 (AH)'},
            {'key': 'ups_max_load_a', 'label': 'UPS 최대부하량 (A)'},
            {'key': 'ups_avg_hr', 'label': 'UPS 평균 (HR)'},
        ],
    },
    {
        'code': '18',
        'title': '옥외설비 설치 기준(신뢰성 고시 별표1)',
        'checks': [
            '풍해 대책 강구 여부',
            "지진 대책 강구 여부('09.10.16. 이후 설치 설비, 임차국사는 '15.7.1.부터 적용)",
            '수해 방지 조치 강구 여부',
            '맨홀 등 내부에 설치되는 접속함체 등의 거치대 이용 설치 여부',
            '지하 통신구의 개구부 설치 위치가 적정하거나, 침수방지장치 설치 여부',
            '제3자의 접촉방지를 위한 맨홀 잠금장치 설치 여부',
            '낙뢰 대책 강구 여부(권고)',
            '진동 대책 강구 여부(권고)',
            '화재 대책 강구 여부(권고)',
            '내수(방수) 기능 구비 여부(권고)',
            '동결 대책 강구 여부(권고)',
            '염해 등 대책(권고)',
            '고온·저온 대책(권고)',
            '다습도 대책(권고)',
            '고신뢰도 부품 사용(해저 등 특수장소)(권고)',
            '제3자의 접촉 및 침입 방지 조치 강구 여부(권고)',
        ],
        'detail_fields': [
            {'key': 'route_tongsingu', 'label': '통신구'},
            {'key': 'route_manhole', 'label': '맨홀'},
            {'key': 'route_gagong', 'label': '가공'},
        ],
    },
    {
        'code': '21',
        'title': '통신기계실의 구조 조건(신뢰성 고시 별표1)',
        'checks': [
            '통신기계실의 전용 공간 설치 여부',
            "바닥시설의 지진대책 기준 적합 여부('09.10.16. 이후 설치 설비, 임차국사는 '15.7.1.부터 적용)",
            '비, 바람, 자외선 및 대기먼지 등에 의한 피해 방지 조치 여부',
            '바닥, 내벽, 천장 내장재의 지진 대비 조치 강구 여부(임차국사)(권고)',
        ],
        'detail_fields': [],
    },
    {
        'code': '22',
        'title': '통신국사 및 통신기계실의 출입제한 기능(신뢰성 고시 별표1)',
        'checks': [
            '모든 출입구에 시건장치 설치 및 출입통제관리 실시 여부(권고)',
        ],
        'detail_fields': [
            {'key': 'restricted_sign', 'label': '제한구역 안내표시 유/무'},
            {'key': 'lock_device', 'label': '시건장치 유/무'},
            {'key': 'cctv', 'label': 'CCTV 운영 유/무'},
        ],
    },
    {
        'code': '23',
        'title': '통신국사의 화재 및 수해대책(신뢰성 고시 별표1)',
        'checks': [
            '주요시설에 소화기, 자동화재탐지설비 및 자동소화설비 설치 여부',
            '주요시설 내장재의 불연재료/준불연재료 여부',
            '통신국사 출입구가 침수 예상 높이보다 높게 설치 여부',
            '통신국사 내 주요시설 지상 설치 여부(또는 방수 조치)',
            '지하공간 출입구 침수 방지턱 설치 여부',
        ],
        'detail_fields': [
            {'key': 'halon_extinguisher', 'label': '하론소화기'},
            {'key': 'powder_extinguisher', 'label': '분말소화기'},
            {'key': 'clean_gas', 'label': '청정가스'},
            {'key': 'cabinet_halon1301', 'label': '케비넷하론1301'},
            {'key': 'centralized_halon', 'label': '집합형 하론'},
            {'key': 'diffusion_extinguisher', 'label': '확산소화기'},
            {'key': 'smoke_detector_ea', 'label': '연기 감지기 (EA)'},
            {'key': 'heat_detector_ea', 'label': '열 감지기 (EA)'},
            {'key': 'differential_detector_ea', 'label': '차동 감지기 (EA)'},
        ],
    },
    {
        'code': '24',
        'title': '통신기계실의 온·습도 관리(신뢰성 고시 별표1)',
        'checks': [
            '항온·항습 기능 구비 여부(권고)',
        ],
        'detail_fields': [
            {'key': 'cooler_count', 'label': '냉방기 (개)'},
        ],
    },
    {
        'code': '25',
        'title': '통신기계실의 분진·유해가스 관리(신뢰성 고시 별표1)',
        'checks': [
            '부식성 가스(SO2 등)나 분진 혼입 시 배제 기능 구비 여부(권고)',
        ],
        'detail_fields': [
            {'key': 'cooling_clean_count', 'label': '냉방기 관리 횟수'},
            {'key': 'equipment_clean_count', 'label': '장비 관리 횟수'},
        ],
    },
    {
        'code': '26',
        'title': '통신망 보전·운용 기준 설정(신뢰성 고시 별표1)',
        'checks': [
            '통신망 보전·운용 기준 설정 및 데이터 집계·관리 여부',
        ],
        'detail_fields': [],
    },
    {
        'code': '31',
        'title': '방송통신설비의 시험 및 결과 기록·관리(방발법 제28조제2항)',
        'checks': [
            '사업자의 방송통신설비 자체 시험 및 기록·관리 여부',
        ],
        'detail_fields': [],
    },
    {
        'code': '32',
        'title': '전송설비 및 선로설비의 보호(규정 제8조 및 접지설비 고시)',
        'checks': [
            '타 설비/차량 통행에 피해가 없도록 설치 여부',
            '하천 횡단 시 안전표지(항공표지 등) 설치 여부',
        ],
        'detail_fields': [],
    },
    {
        'code': '33',
        'title': '선로설비의 설치 및 철거방법(규정 제18조 및 접지설비 고시)',
        'checks': [
            '가공통신선 지지물의 등주 방지 적정 여부',
            '가공통신선의 높이 적정 여부',
            '통신선과 전력선 간 이격거리 적정 여부',
            '동일 지지물에서 통신선/전력선 이격거리 적정 여부',
            '지중/해저 통신선과 전력선 이격거리 적정 여부',
        ],
        'detail_fields': [],
    },
    {
        'code': '34',
        'title': '통신장비류의 지진 대책(신뢰성 고시 제5조 및 별표2)',
        'checks': [
            '통신장비, 전원설비, 부대설비의 지진 대책 적정 여부',
        ],
        'detail_fields': [],
    },
]

FACILITY_CAPTURE_SLOTS = [
    {
        'slot_no': 1,
        'title': '보호기 설치여부',
        'detail': '인입단에 서지보호기 설치'
    },
    {
        'slot_no': 2,
        'title': '접지저항 적합여부',
        'detail': '메인 접지저항(00 Ω)'
    },
    {
        'slot_no': 3,
        'title': '예비전원 설비',
        'detail': '축전지 용량 및 배선 적합여부'
    },
    {
        'slot_no': 4,
        'title': '제3자의 접촉 방지',
        'detail': '자물쇠, 시설 보호망(펜스) 등 설치'
    },
    {
        'slot_no': 5,
        'title': '통신설비 환경 저해요인 제거',
        'detail': '통신국사 까치집 제거 및 잡초제거 여부'
    },
    {
        'slot_no': 6,
        'title': '내진대상설비 지진대책',
        'detail': '통신장비 고정앙카 또는 볼트 설치'
    },
    {
        'slot_no': 7,
        'title': '옥외설비 육안검사',
        'detail': '통신주 볼트의 풀림 및 파손여부 등'
    },
    {
        'slot_no': 8,
        'title': '통신설비 화재대책',
        'detail': '소화기, 감지장치, 자동소화장치 등'
    },
]


def build_default_facility_payload():
    sections = {}
    for section in FACILITY_SURVEY_SECTIONS:
        sections[section['code']] = {
            'result': '',
            'checks': ['' for _ in section.get('checks', [])],
            'detail_note': '',
            'fields': {field['key']: '' for field in section.get('detail_fields', [])}
        }
    return {
        'version': FACILITY_SURVEY_VERSION,
        'business_name': '',
        'site_type': '',
        'inspection_date': '',
        'sections': sections
    }


def parse_facility_payload(raw_checklist_data):
    payload = build_default_facility_payload()
    if not raw_checklist_data:
        return payload

    try:
        parsed = raw_checklist_data if isinstance(raw_checklist_data, dict) else json.loads(raw_checklist_data)
    except Exception:
        return payload

    if not isinstance(parsed, dict):
        return payload

    payload['business_name'] = str(parsed.get('business_name', '') or '').strip()
    payload['site_type'] = str(parsed.get('site_type', '') or '').strip()
    payload['inspection_date'] = str(parsed.get('inspection_date', '') or '').strip()

    source_sections = parsed.get('sections') or {}
    if not isinstance(source_sections, dict):
        return payload

    for section in FACILITY_SURVEY_SECTIONS:
        code = section['code']
        src = source_sections.get(code) or {}
        if not isinstance(src, dict):
            continue

        result = str(src.get('result', '') or '').strip()
        if result in FACILITY_RESULT_OPTIONS:
            payload['sections'][code]['result'] = result

        checks = src.get('checks') or []
        if isinstance(checks, list):
            safe_checks = []
            for idx, _check_name in enumerate(section.get('checks', [])):
                value = str(checks[idx] if idx < len(checks) else '').strip()
                safe_checks.append(value if value in FACILITY_CHECK_OPTIONS else '')
            payload['sections'][code]['checks'] = safe_checks

        payload['sections'][code]['detail_note'] = str(src.get('detail_note', '') or '').strip()

        src_fields = src.get('fields') or {}
        if isinstance(src_fields, dict):
            for field in section.get('detail_fields', []):
                key = field['key']
                payload['sections'][code]['fields'][key] = str(src_fields.get(key, '') or '').strip()

    return payload


def build_facility_payload_from_form(form_data):
    payload = build_default_facility_payload()
    payload['business_name'] = form_data.get('business_name', '').strip()
    payload['site_type'] = form_data.get('site_type', '').strip()
    payload['inspection_date'] = form_data.get('inspection_date', '').strip()

    for section in FACILITY_SURVEY_SECTIONS:
        code = section['code']
        section_state = payload['sections'][code]

        result = form_data.get(f'sec_{code}_result', '').strip()
        section_state['result'] = result if result in FACILITY_RESULT_OPTIONS else ''

        checks = []
        for idx, _check_name in enumerate(section.get('checks', []), start=1):
            value = form_data.get(f'sec_{code}_check_{idx}', '').strip()
            checks.append(value if value in FACILITY_CHECK_OPTIONS else '')
        section_state['checks'] = checks

        section_state['detail_note'] = form_data.get(f'sec_{code}_detail_note', '').strip()

        for field in section.get('detail_fields', []):
            key = field['key']
            section_state['fields'][key] = form_data.get(f'sec_{code}_field_{key}', '').strip()

    return payload


def format_equipment_no(index, prefix=DEFAULT_EQUIPMENT_PREFIX):
    safe_index = max(1, int(index))
    return f"{prefix}-{safe_index:02d}"


def parse_equipment_no(equipment_no):
    if not equipment_no:
        return None
    normalized = str(equipment_no).strip().upper()
    match = EQUIPMENT_NO_PATTERN.match(normalized)
    if not match:
        return None
    return match.group('prefix'), int(match.group('index'))


def sort_equipment_nos(equipment_numbers):
    parsed_rows = []
    for raw_no in equipment_numbers or []:
        parsed = parse_equipment_no(raw_no)
        if parsed:
            parsed_rows.append((parsed[0], parsed[1], raw_no))
        else:
            parsed_rows.append(('ZZZ', 9999, raw_no))
    parsed_rows.sort(key=lambda row: (row[0], row[1]))
    return [row[2] for row in parsed_rows]


def get_next_equipment_no(equipment_numbers):
    max_index = 0
    for raw_no in equipment_numbers or []:
        parsed = parse_equipment_no(raw_no)
        if parsed and parsed[0] == DEFAULT_EQUIPMENT_PREFIX:
            max_index = max(max_index, parsed[1])
    return format_equipment_no(max_index + 1 if max_index else DEFAULT_EQUIPMENT_START)


INSPECTION_ITEMS = [
    (1, '고무패킹교체'),
    (2, '실내기 Reset'),
    (3, 'V벨트 교체'),
    (4, '타이머 릴레이'),
    (5, '배수관 청소'),
    (6, 'RMS 온도센싱'),
    (7, '자연공조 점검'),
    (8, '정전보상'),
    (9, '실외기 핀,넝쿨'),
    (10, '송풍구 풍량'),
    (11, '열화상 측정'),
]

INSPECTION_METHOD_GUIDE = {
    1: "공냉식 토출구가 그릴타입 대상\n- 위치고정이 안되는 경우\n- 고무패킹 교체",
    2: "공냉식 micom 램프 확인 후 Reset\n- RUN 램프 미점등: ON/OFF버튼 눌러 RUN 점등 확인\n- Alarm 램프 점등: Reset 버튼 누름 후 램프 복구 확인",
    3: "공냉식(구형) 공기 토출구에서 V벨트 존재 확인\n- 앞판을 열어 벨트 손상/처짐 여부 확인",
    4: "공냉식 릴레이, 타이머 불량 점검\n- 불량 시 BP 교체요청\n- 변색 또는 열화상 측정 시 고온 발생 여부 확인",
    5: "(전체) 오염물질 많으면 청소\n- 전후 사진\n- 실내기 필터\n- 배수구 청소 등",
    6: "(전체) 온도센서 설치위치 점검\n- 이상한 경우(고온) 확인",
    7: "집중국 기반 외부공기 팬/모터 점검",
    8: "공냉식 micom 내부 분리 후 dip s/w 확인 필요\n- 좌측에 있으면 우측으로 변경",
    9: "(전체) 청소 시 전후사진\n- 칡넝쿨 등 조치 시",
    10: "(전체) 풍량계 측정 또는 촉감점검",
    11: "(전체) 분전반 커버 탈착 후 열화상 측정\n- 사진 보관",
}

FACILITY_INSPECTION_SITE_NAMES = [
    "대림2WD",
    "대방역2WD",
]

INSPECTION_SITE_NAMES = [
    "(HK)수서역LDT1.51.LTE.DU30(내)",
    "(RM)천호2LDB.51.LTE.ENB(내)",
    "BR_여의도W.51.WCDMA.E3-NODEB(내)",
    "BR_외수협W.51.WCDMA.E3-NODEB(내)",
    "GR_강서구청W.51.WCDMA.E3-NODEB(내)",
    "GR_내곡동INFOB7LDC_00.51.LTE.DU20(내)",
    "GR_봉천4동2W.51.WCDMA.E3-NODEB(외)",
    "GR_의사당WT1.51.WCDMA.E3-NODEB(내)-T",
    "PR_목6동W.51.WCDMA.E3-NODEB(내)",
    "SKB동작사옥W.51.WCDMA.E3-NODEB(외)",
    "가락W.51.WCDMA.E3-NODEB(내)",
    "가리봉(1FA)WD1.51.WCDMA.IPNB-DU-6S(내)",
    "가산동2(1FA)WD1.51.WCDMA.IPNB-DU-3S(외)",
    "가양(1FA)WD1.51.WCDMA.IPNB-DU-3S(내)",
    "공항(1FA)WD5.51.WCDMA.IPNB-DU-6S(내)",
    "관악소방서4TLDT_00.51.LTE.DU30(내)",
    "광명대교(1FA)WD.51.WCDMA.IPNB-DU-9S(내)",
    "구로에이스WT1.51.WCDMA.E3-NODEB(외)",
    "구로역W.51.WCDMA.E3-NODEB(내)",
    "궁동2(1FA)WD2.51.WCDMA.IPNB-DU-3S(내)",
    "길동(1FA)WD.51.WCDMA.IPNB-DU-9S(내)",
    "낙성대(1FA)WD.51.WCDMA.IPNB-DU-6S(내)",
    "내곡IC_LDT_00.51.LTE.DU30(내)",
    "노량진(1FA)WD3.51.WCDMA.IPNB-DU-6S(내)",
    "당산3W.51.WCDMA.E3-NODEB(내)",
    "대림2(1FA)WD.51.WCDMA.IPNB-DU-6S(내)",
    "대림W.51.WCDMA.E3-NODEB(내)",
    "대림우체국W.51.WCDMA.E3-NODEB(내)",
    "대치(1FA)WD1.51.WCDMA.IPNB-DU-9S(내)",
    "도곡주공(1FA)WD1.51.WCDMA.IPNB-DU-3S(내)",
    "도곡행자타운(1FA)WD2.51.WCDMA.IPNB-DU-3S(내)",
    "독산2(1FA)WD2.51.WCDMA.IPNB-DU-3S(내)",
    "명덕고교(1FA)WD1.51.WCDMA.IPNB-DU-3S(내)",
    "반포W.51.WCDMA.E3-NODEB(내)",
    "방배사거리W.51.WCDMA.E3-NODEB(외)",
    "상도(1FA)WD4.51.WCDMA.IPNB-DU-6S(내)",
    "서초3_4TLDT_00.51.LTE.DU30(내)",
    "수서원(1FA)WD1.51.WCDMA.IPNB-DU-3S(내)",
    "시흥SE(1FA)WD1.51.WCDMA.IPNB-DU-3S(내)",
    "신림2(1FA)WD.51.WCDMA.IPNB-DU-6S(내)",
    "신림7동(1FA)WD1.51.WCDMA.IPNB-DU-6S(내)",
    "신림분동(1FA)WD.51.WCDMA.IPNB-DU-6S(내)",
    "신사(1FA)WD1.51.WCDMA.IPNB-DU-6S(내)",
    "여의하류B7LDC_00.51.LTE.DU25(내)",
    "영등6교(1FA)WD1.51.WCDMA.IPNB-DU-3S(외)",
    "오류ICW.51.WCDMA.E3-NODEB(외)",
    "오류동W.51.WCDMA.E3-NODEB(외)",
    "일원역(1FA)WD.51.WCDMA.IPNB-DU-9S(내)",
    "포스코사W.51.WCDMA.E3-NODEB(내)",
    "풍납(1FA)WD1.51.WCDMA.IPNB-DU-3S(내)",
    "한강대교2(1FA)WD1.51.WCDMA.IPNB-DU-3S(내)",
    "한강대교WT.51.WCDMA.E3-NODEB(내)",
    "화곡2W.51.WCDMA.E3-NODEB(내)",
]

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_korea_time():
    korea_tz = pytz.timezone('Asia/Seoul')
    return datetime.now(korea_tz)

def find_local_inspection_method_image_relpath():
    inspection_dir = os.path.join(app.root_path, 'static', 'inspection')
    os.makedirs(inspection_dir, exist_ok=True)

    candidate_files = [
        'inspection_method.png',
        'inspection_method.jpg',
        'inspection_method.jpeg',
        'inspection_method.webp',
        'inspection-method.png',
        'inspection-method.jpg',
        'inspection-method.jpeg',
        'inspection-method.webp'
    ]

    for filename in candidate_files:
        abs_path = os.path.join(inspection_dir, filename)
        if os.path.exists(abs_path):
            return f'inspection/{filename}'

    for filename in os.listdir(inspection_dir):
        lower_name = filename.lower()
        if lower_name.endswith(('.png', '.jpg', '.jpeg', '.webp')):
            return f'inspection/{filename}'

    return None

def get_latest_inspection_method_image():
    """점검방법 이미지(DB+Supabase) 최신 활성 데이터 조회"""
    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
            '''SELECT id, storage_path, public_url, uploaded_by, created_at
               FROM inspection_method_images
               WHERE is_active = 1
               ORDER BY created_at DESC, id DESC
               LIMIT 1'''
        )
        row = cursor.fetchone()
        conn.close()
        if not row:
            return None
        return {
            'id': row[0],
            'storage_path': row[1],
            'public_url': row[2],
            'uploaded_by': row[3],
            'created_at': row[4]
        }
    except Exception:
        try:
            if conn:
                conn.close()
        except Exception:
            pass
        return None

def fetch_image_bytes(image_url):
    if not image_url:
        return None
    try:
        response = requests.get(image_url, timeout=20)
        if response.status_code != 200:
            return None
        content_type = response.headers.get('content-type', '').lower()
        if 'image' not in content_type and not response.content:
            return None
        return response.content
    except Exception:
        return None


def safe_excel_sheet_title(title, fallback='Sheet1'):
    invalid_chars = set('[]:*?/\\')
    cleaned = ''.join(ch for ch in str(title or '') if ch not in invalid_chars).strip()
    if not cleaned:
        cleaned = fallback
    return cleaned[:31]

def excel_column_width_to_pixels(width):
    """Excel 열 너비 값을 픽셀로 근사 변환."""
    if not width:
        width = 8.43
    return int(width * 7 + 5)


def excel_row_height_to_pixels(height):
    """Excel 행 높이(pt)를 픽셀로 근사 변환."""
    if not height:
        height = 15
    return int(height * 96 / 72)


def add_excel_image(worksheet, row_no, col_no, image_bytes, image_refs, max_width=None, max_height=None):
    if not image_bytes:
        return False
    try:
        pil_image = Image.open(io.BytesIO(image_bytes))
        if pil_image.mode not in ('RGB', 'RGBA'):
            pil_image = pil_image.convert('RGB')

        width, height = pil_image.size
        if width <= 0 or height <= 0:
            return False

        col_letter = get_column_letter(col_no)
        column_width = worksheet.column_dimensions[col_letter].width
        row_height = worksheet.row_dimensions[row_no].height

        target_width = int(max_width) if max_width else excel_column_width_to_pixels(column_width)
        target_height = int(max_height) if max_height else excel_row_height_to_pixels(row_height)
        target_width = max(target_width, 1)
        target_height = max(target_height, 1)

        # 셀에 꽉 차도록 이미지 비율 유지 대신 셀 크기로 맞춤.
        resized = pil_image.resize((target_width, target_height), Image.Resampling.LANCZOS)

        if resized.mode != 'RGB':
            resized = resized.convert('RGB')

        output = io.BytesIO()
        resized.save(output, format='JPEG', quality=85, optimize=True)
        output.seek(0)

        excel_image = XLImage(output)
        excel_image.width = target_width
        excel_image.height = target_height
        excel_image.anchor = f"{get_column_letter(col_no)}{row_no}"
        worksheet.add_image(excel_image)
        image_refs.append(output)
        return True
    except Exception:
        return False

def get_db_connection():
    """안정적인 데이터베이스 연결 함수"""
    try:
        import pg8000
        parsed = urllib.parse.urlparse(DATABASE_URL)
        
        conn = pg8000.connect(
            host=parsed.hostname,
            port=parsed.port or 5432,
            user=parsed.username,
            password=parsed.password,
            database=parsed.path[1:] if parsed.path else 'postgres'
        )
        
        conn.autocommit = False
        
        return conn
    except ImportError:
        print("❌ 치명적 오류: pg8000 라이브러리가 설치되지 않았습니다!")
        raise Exception("pg8000 라이브러리 필요")
    except Exception as e:
        print(f"❌ 치명적 오류: Supabase PostgreSQL 연결 실패!")
        print(f"   오류 내용: {e}")
        raise Exception(f"Supabase 연결 실패: {e}")

def send_email(to_emails, subject, html_content):
    """이메일 발송 함수"""
    try:
        if not SMTP_USERNAME or not SMTP_PASSWORD:
            return False, "이메일 설정이 되어있지 않습니다."
        
        msg = MIMEMultipart('alternative')
        msg['From'] = SMTP_USERNAME
        msg['To'] = ', '.join(to_emails) if isinstance(to_emails, list) else to_emails
        msg['Subject'] = subject
        
        html_part = MIMEText(html_content, 'html', 'utf-8')
        msg.attach(html_part)
        
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SMTP_USERNAME, SMTP_PASSWORD)
        
        text = msg.as_string()
        server.sendmail(SMTP_USERNAME, to_emails, text)
        server.quit()
        
        return True, "이메일이 성공적으로 발송되었습니다."
        
    except Exception as e:
        print(f"이메일 발송 오류: {e}")
        return False, f"이메일 발송 실패: {str(e)}"

def compress_image_to_target_size(image_file, max_size_mb=1, max_width=800, quality=85):
    """
    이미지를 목표 크기(MB) 이하로 압축하는 함수
    
    Args:
        image_file: 업로드된 이미지 파일
        max_size_mb: 최대 파일 크기 (MB)
        max_width: 최대 가로 크기 (픽셀)
        quality: JPEG 품질 (20-95)
    
    Returns:
        compressed_image_bytes: 압축된 이미지 바이트
        final_size_kb: 최종 파일 크기 (KB)
    """
    try:
        # PIL Image로 열기
        img = Image.open(image_file)
        
        # EXIF 회전 정보 처리 (스마트폰 사진)
        if hasattr(img, '_getexif') and img._getexif() is not None:
            exif = img._getexif()
            orientation = exif.get(274)
            if orientation == 3:
                img = img.rotate(180, expand=True)
            elif orientation == 6:
                img = img.rotate(270, expand=True)
            elif orientation == 8:
                img = img.rotate(90, expand=True)
        
        # RGB 모드로 변환 (JPEG 저장용)
        if img.mode in ('RGBA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = background
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        
        # 원본 크기 계산
        original_width, original_height = img.size
        
        # 크기 조정 (비율 유지)
        if original_width > max_width:
            ratio = max_width / original_width
            new_height = int(original_height * ratio)
            img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
        
        # 목표 크기까지 품질 조정하면서 압축
        max_size_bytes = max_size_mb * 1024 * 1024
        current_quality = quality
        
        while current_quality > 20:
            output = io.BytesIO()
            img.save(output, format='JPEG', quality=current_quality, optimize=True)
            
            if output.tell() <= max_size_bytes:
                break
                
            current_quality -= 5
            output.seek(0)
            output.truncate(0)
        
        output.seek(0)
        compressed_bytes = output.getvalue()
        final_size_kb = len(compressed_bytes) / 1024
        
        print(f"✅ 이미지 압축 완료: {final_size_kb:.1f}KB (품질: {current_quality})")
        
        return compressed_bytes, final_size_kb
        
    except Exception as e:
        print(f"❌ 이미지 압축 오류: {e}")
        return None, 0

def upload_to_supabase_storage(image_bytes, filename, bucket_name='warehouse-photos', content_type='image/jpeg'):
    """
    압축된 이미지를 Supabase Storage에 업로드
    
    Args:
        image_bytes: 압축된 이미지 바이트
        filename: 저장할 파일명
        bucket_name: Supabase Storage 버킷명
    
    Returns:
        public_url: 업로드된 파일의 공개 URL
    """
    try:
        # Supabase Storage API 엔드포인트
        upload_url = f"{SUPABASE_URL}/storage/v1/object/{bucket_name}/{filename}"
        
        headers = {
            'Authorization': f'Bearer {SUPABASE_SERVICE_KEY}',
            'Content-Type': content_type
        }
        
        # 파일 업로드
        response = requests.post(upload_url, data=image_bytes, headers=headers)
        
        if response.status_code in [200, 201]:
            # 공개 URL 생성
            public_url = f"{SUPABASE_URL}/storage/v1/object/public/{bucket_name}/{filename}"
            print(f"✅ Supabase Storage 업로드 성공: {public_url}")
            return public_url
        else:
            print(f"❌ Supabase Storage 업로드 실패: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        print(f"❌ Supabase Storage 업로드 오류: {e}")
        return None


def save_inspection_photo(file_obj, item_id, checkpoint_no, phase):
    """점검 사진을 1MB 이하로 압축 후 Supabase에 저장"""
    if not file_obj or file_obj.filename == '':
        raise Exception(f"{checkpoint_no}번 {phase} 사진이 없습니다.")
    if not allowed_file(file_obj.filename):
        raise Exception(f"{checkpoint_no}번 {phase} 사진 형식이 올바르지 않습니다.")

    file_obj.seek(0)
    compressed_bytes, final_size_kb = compress_image_to_target_size(
        file_obj,
        max_size_mb=0.9,
        max_width=1600,
        quality=85
    )
    if not compressed_bytes:
        raise Exception(f"{checkpoint_no}번 {phase} 사진 압축에 실패했습니다.")

    filename = f"inspection_{item_id}_{checkpoint_no}_{phase}_{uuid.uuid4().hex}.jpg"
    supabase_url = upload_to_supabase_storage(compressed_bytes, filename)
    if not supabase_url:
        raise Exception(f"{checkpoint_no}번 {phase} 사진 업로드에 실패했습니다.")

    return filename, int(final_size_kb), supabase_url


def save_facility_capture_photo(file_obj, item_id, slot_no):
    """시설물 적합조사 사진(항목당 1장)을 1MB 미만으로 압축 후 저장"""
    if not file_obj or file_obj.filename == '':
        return None
    if not allowed_file(file_obj.filename):
        raise Exception(f"{slot_no}번 사진 형식이 올바르지 않습니다.")

    file_obj.seek(0)
    compressed_bytes, final_size_kb = compress_image_to_target_size(
        file_obj,
        max_size_mb=0.95,
        max_width=1600,
        quality=88
    )
    if not compressed_bytes:
        raise Exception(f"{slot_no}번 사진 압축에 실패했습니다.")

    filename = f"facility_capture_{item_id}_{slot_no}_{uuid.uuid4().hex}.jpg"
    supabase_url = upload_to_supabase_storage(compressed_bytes, filename)
    if not supabase_url:
        raise Exception(f"{slot_no}번 사진 업로드에 실패했습니다.")

    return {
        'filename': filename,
        'file_size': int(final_size_kb),
        'supabase_url': supabase_url
    }


def fetch_facility_capture_photos(cursor, record_id):
    photo_map = {}
    cursor.execute(
        '''SELECT slot_no, point_title, detail_text, filename, file_size, supabase_url
           FROM facility_capture_photos
           WHERE record_id = %s
           ORDER BY slot_no''',
        (record_id,)
    )
    for slot_no, point_title, detail_text, filename, file_size, supabase_url in cursor.fetchall():
        photo_map[int(slot_no)] = {
            'point_title': point_title or '',
            'detail_text': detail_text or '',
            'filename': filename,
            'file_size': file_size,
            'supabase_url': supabase_url
        }
    return photo_map

def init_db():
    """트랜잭션 오류 완전 해결된 초기화 함수"""
    conn = None
    try:
        print("🔄 Supabase PostgreSQL 연결 테스트 중...")
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('SELECT version()')
        version_info = cursor.fetchone()[0]
        print(f"✅ Supabase 연결 성공!")
        print(f"📊 PostgreSQL 버전: {version_info[:50]}...")
        
        print("🔄 데이터베이스 테이블 생성 중...")
        
        # 각 테이블을 개별 트랜잭션으로 생성
        tables_to_create = [
            ('users', '''CREATE TABLE IF NOT EXISTS users (
                id SERIAL PRIMARY KEY,
                name TEXT NOT NULL,
                employee_id TEXT UNIQUE NOT NULL,
                team TEXT NOT NULL,
                password TEXT NOT NULL,
                is_approved INTEGER DEFAULT 0,
                created_at TIMESTAMP DEFAULT (NOW() AT TIME ZONE 'Asia/Seoul')
            )'''),
            ('inventory', '''CREATE TABLE IF NOT EXISTS inventory (
                id SERIAL PRIMARY KEY,
                warehouse TEXT NOT NULL,
                category TEXT NOT NULL,
                part_name TEXT NOT NULL,
                quantity INTEGER DEFAULT 0,
                last_modifier TEXT,
                last_modified TIMESTAMP DEFAULT (NOW() AT TIME ZONE 'Asia/Seoul')
            )'''),
            ('inventory_history', '''CREATE TABLE IF NOT EXISTS inventory_history (
                id SERIAL PRIMARY KEY,
                inventory_id INTEGER REFERENCES inventory(id),
                change_type TEXT,
                quantity_change INTEGER,
                modifier_name TEXT,
                modified_at TIMESTAMP DEFAULT (NOW() AT TIME ZONE 'Asia/Seoul')
            )'''),
            ('photos', '''CREATE TABLE IF NOT EXISTS photos (
                id SERIAL PRIMARY KEY,
                inventory_id INTEGER REFERENCES inventory(id),
                filename TEXT NOT NULL,
                original_name TEXT NOT NULL,
                file_size INTEGER,
                uploaded_by TEXT,
                uploaded_at TIMESTAMP DEFAULT (NOW() AT TIME ZONE 'Asia/Seoul'),
                supabase_url TEXT
            )'''),
            ('delivery_receipts', '''CREATE TABLE IF NOT EXISTS delivery_receipts (
                id SERIAL PRIMARY KEY,
                receipt_date DATE NOT NULL,
                receipt_type TEXT NOT NULL,
                items_data TEXT,
                signature_data TEXT,
                created_by TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT (NOW() AT TIME ZONE 'Asia/Seoul')
            )'''),
            ('inspection_records', '''CREATE TABLE IF NOT EXISTS inspection_records (
                id SERIAL PRIMARY KEY,
                inventory_id INTEGER REFERENCES inventory(id),
                warehouse TEXT NOT NULL,
                site_name TEXT,
                equipment_no TEXT,
                inspector_name TEXT NOT NULL,
                inspected_at TIMESTAMP DEFAULT (NOW() AT TIME ZONE 'Asia/Seoul'),
                status TEXT DEFAULT '작업 완료',
                checklist_data TEXT,
                memo TEXT
            )'''),
            ('inspection_units', '''CREATE TABLE IF NOT EXISTS inspection_units (
                id SERIAL PRIMARY KEY,
                inventory_id INTEGER REFERENCES inventory(id) ON DELETE CASCADE,
                equipment_no TEXT NOT NULL,
                created_by TEXT,
                created_at TIMESTAMP DEFAULT (NOW() AT TIME ZONE 'Asia/Seoul'),
                UNIQUE(inventory_id, equipment_no)
            )'''),
            ('inspection_photos', '''CREATE TABLE IF NOT EXISTS inspection_photos (
                id SERIAL PRIMARY KEY,
                record_id INTEGER REFERENCES inspection_records(id) ON DELETE CASCADE,
                checkpoint_no INTEGER NOT NULL,
                checkpoint_name TEXT NOT NULL,
                phase TEXT NOT NULL,
                filename TEXT NOT NULL,
                file_size INTEGER,
                supabase_url TEXT NOT NULL,
                uploaded_at TIMESTAMP DEFAULT (NOW() AT TIME ZONE 'Asia/Seoul')
            )'''),
            ('facility_capture_photos', '''CREATE TABLE IF NOT EXISTS facility_capture_photos (
                id SERIAL PRIMARY KEY,
                record_id INTEGER REFERENCES inspection_records(id) ON DELETE CASCADE,
                slot_no INTEGER NOT NULL,
                point_title TEXT,
                detail_text TEXT,
                filename TEXT NOT NULL,
                file_size INTEGER,
                supabase_url TEXT NOT NULL,
                uploaded_at TIMESTAMP DEFAULT (NOW() AT TIME ZONE 'Asia/Seoul'),
                UNIQUE(record_id, slot_no)
            )'''),
            ('inspection_method_images', '''CREATE TABLE IF NOT EXISTS inspection_method_images (
                id SERIAL PRIMARY KEY,
                storage_path TEXT NOT NULL,
                public_url TEXT NOT NULL,
                uploaded_by TEXT,
                created_at TIMESTAMP DEFAULT (NOW() AT TIME ZONE 'Asia/Seoul'),
                is_active INTEGER DEFAULT 1
            )'''),
            ('app_settings', '''CREATE TABLE IF NOT EXISTS app_settings (
                setting_key TEXT PRIMARY KEY,
                setting_value TEXT,
                updated_at TIMESTAMP DEFAULT (NOW() AT TIME ZONE 'Asia/Seoul')
            )''')
        ]
        
        for table_name, sql in tables_to_create:
            try:
                cursor.execute(sql)
                conn.commit()
                print(f"✅ {table_name} 테이블 처리 완료")
            except Exception as e:
                conn.rollback()
                print(f"⚠️ {table_name} 테이블 처리 중 오류 (무시): {e}")
                cursor.close()
                cursor = conn.cursor()
        
        # supabase_url 컬럼 추가 (이미 존재할 수 있으므로 오류 무시)
        try:
            cursor.execute('ALTER TABLE photos ADD COLUMN supabase_url TEXT')
            conn.commit()
            print("✅ photos 테이블에 supabase_url 컬럼 추가 완료")
        except Exception as e:
            conn.rollback()
            print(f"ℹ️ supabase_url 컬럼 이미 존재 또는 추가 불필요: {e}")
            cursor.close()
            cursor = conn.cursor()
        
        # 관리자 계정 생성 (별도 트랜잭션)
        try:
            cursor.execute('ALTER TABLE inspection_records ADD COLUMN equipment_no TEXT')
            conn.commit()
            print("??inspection_records ?뚯씠釉붿뿉 equipment_no 而щ읆 異붽? ?꾨즺")
        except Exception as e:
            conn.rollback()
            print(f"?뱄툘 equipment_no 而щ읆 ?대? 議댁옱 ?먮뒗 異붽? 遺덊븘?? {e}")
            cursor.close()
            cursor = conn.cursor()

        try:
            cursor.execute('SELECT id FROM users WHERE employee_id = %s', ('admin',))
            admin_exists = cursor.fetchone()
            
            if not admin_exists:
                admin_password_hash = generate_password_hash(ADMIN_PASSWORD)
                cursor.execute('''INSERT INTO users (name, employee_id, team, password, is_approved) 
                                 VALUES (%s, %s, %s, %s, %s)''',
                              ('관리자', 'admin', '관리', admin_password_hash, 1))
                conn.commit()
                print("✅ 관리자 계정 생성 완료")
            else:
                print("ℹ️ 관리자 계정 이미 존재")
                
        except Exception as admin_error:
            conn.rollback()
            print(f"⚠️ 관리자 계정 처리 중 오류: {admin_error}")

        # 점검 대상(국사명) 시드 데이터 추가 (최초 1회만)
        try:
            cursor.execute(
                "SELECT setting_value FROM app_settings WHERE setting_key = %s",
                ('inspection_seed_initialized',)
            )
            seed_flag = cursor.fetchone()

            if seed_flag:
                print("ℹ️ 점검 대상 시드 이미 초기화됨 - 자동 재삽입 건너뜀")
            else:
                cursor.execute(
                    '''SELECT COUNT(*)
                       FROM inventory
                       WHERE warehouse = %s AND category = %s''',
                    (DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY)
                )
                existing_diy_count = cursor.fetchone()[0] or 0

                if existing_diy_count > 0:
                    cursor.execute(
                        '''INSERT INTO app_settings (setting_key, setting_value)
                           VALUES (%s, %s)
                           ON CONFLICT (setting_key)
                           DO UPDATE SET setting_value = EXCLUDED.setting_value, updated_at = (NOW() AT TIME ZONE 'Asia/Seoul')''',
                        ('inspection_seed_initialized', 'existing-data')
                    )
                    conn.commit()
                    print(f"ℹ️ 기존 점검 대상 {existing_diy_count}건 확인 - 시드 초기화 플래그만 저장")
                else:
                    inserted_count = 0
                    migrated_count = 0
                    for site_name in INSPECTION_SITE_NAMES:
                        cursor.execute(
                            '''SELECT id
                               FROM inventory
                               WHERE warehouse = %s AND category = %s AND part_name = %s''',
                            (DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY, site_name)
                        )
                        if cursor.fetchone():
                            continue

                        cursor.execute(
                            '''SELECT id
                               FROM inventory
                               WHERE warehouse = %s AND category = %s AND part_name = %s
                               ORDER BY id
                               LIMIT 1''',
                            (DB_ACTIVE_WAREHOUSE, '전기차', site_name)
                        )
                        legacy_row = cursor.fetchone()
                        if legacy_row:
                            cursor.execute(
                                '''UPDATE inventory
                                   SET category = %s
                                   WHERE id = %s''',
                                (DIY_CHECKLIST_CATEGORY, legacy_row[0])
                            )
                            migrated_count += 1
                            continue

                        cursor.execute(
                            '''INSERT INTO inventory
                               (warehouse, category, part_name, quantity, last_modifier)
                               VALUES (%s, %s, %s, %s, %s)''',
                            (DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY, site_name, 0, "system-seed")
                        )
                        inserted_count += 1

                    cursor.execute(
                        '''INSERT INTO app_settings (setting_key, setting_value)
                           VALUES (%s, %s)
                           ON CONFLICT (setting_key)
                           DO UPDATE SET setting_value = EXCLUDED.setting_value, updated_at = (NOW() AT TIME ZONE 'Asia/Seoul')''',
                        ('inspection_seed_initialized', f'seeded:{inserted_count}|migrated:{migrated_count}')
                    )
                    conn.commit()
                    print(f"✅ 점검 대상 시드 반영 완료 (신규 {inserted_count}건, 전환 {migrated_count}건)")
        except Exception as seed_error:
            conn.rollback()
            print(f"⚠️ 점검 대상 시드 반영 중 오류: {seed_error}")

        try:
            inserted_count = 0
            for site_name in FACILITY_INSPECTION_SITE_NAMES:
                cursor.execute(
                    '''SELECT id
                       FROM inventory
                       WHERE warehouse = %s AND category = %s AND part_name = %s''',
                    (DB_FACILITY_WAREHOUSE, DIY_CHECKLIST_CATEGORY, site_name)
                )
                if cursor.fetchone():
                    continue

                cursor.execute(
                    '''INSERT INTO inventory
                       (warehouse, category, part_name, quantity, last_modifier)
                       VALUES (%s, %s, %s, %s, %s)''',
                    (DB_FACILITY_WAREHOUSE, DIY_CHECKLIST_CATEGORY, site_name, 0, "system-seed")
                )
                inserted_count += 1

            conn.commit()
            print(f"✅ 시설물 적합조사 점검 대상 시드 반영 완료 (신규 {inserted_count}건)")
        except Exception as facility_seed_error:
            conn.rollback()
            print(f"⚠️ 시설물 적합조사 점검 대상 시드 반영 중 오류: {facility_seed_error}")

        try:
            cursor.execute(
                '''SELECT DISTINCT r.inventory_id
                   FROM inspection_records r
                   JOIN inventory i ON i.id = r.inventory_id
                   WHERE i.warehouse = %s
                     AND i.category = %s''',
                (DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY)
            )
            legacy_inventory_ids = [row[0] for row in cursor.fetchall()]
            default_equipment_no = format_equipment_no(DEFAULT_EQUIPMENT_START)

            for inventory_id in legacy_inventory_ids:
                cursor.execute(
                    '''INSERT INTO inspection_units (inventory_id, equipment_no, created_by)
                       VALUES (%s, %s, %s)
                       ON CONFLICT (inventory_id, equipment_no) DO NOTHING''',
                    (inventory_id, default_equipment_no, 'system-migration')
                )

            cursor.execute(
                '''UPDATE inspection_records r
                   SET equipment_no = %s
                   FROM inventory i
                   WHERE r.inventory_id = i.id
                     AND i.warehouse = %s
                     AND i.category = %s
                     AND (r.equipment_no IS NULL OR TRIM(r.equipment_no) = '')''',
                (default_equipment_no, DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY)
            )
            conn.commit()
        except Exception as migration_error:
            conn.rollback()
            print(f"⚠️ 설비번호 마이그레이션 중 오류: {migration_error}")

    except Exception as e:
        if conn:
            conn.rollback()
        print(f"❌ 초기화 중 오류: {e}")
        raise
    finally:
        if conn:
            conn.close()
        print("✅ 데이터베이스 초기화 완료!")

# 시스템 시작 시 Supabase 연결 필수 확인
print("🔍 Supabase 연결 상태 확인 중...")
init_db()
print("=" * 60)
print("✅ 시스템 준비 완료 - Supabase 연결됨")
print("=" * 60)

# ========
# 디버깅용 함수
# ========
def log_session_debug(route_name):
    """세션 디버깅 로그"""
    print(f"🔍 [{route_name}] 세션 상태:")
    print(f"   user_id: {session.get('user_id', 'None')}")
    print(f"   is_admin: {session.get('is_admin', 'None')}")
    print(f"   user_name: {session.get('user_name', 'None')}")
    print(f"   세션 키들: {list(session.keys())}")

# ========
# 기존 라우트들 (변경사항 없음)
# ========
@app.route('/')
def index():
    """메인 페이지 - 로그인된 사용자는 적절한 대시보드로 리다이렉트"""
    log_session_debug('/')
    
    if 'user_id' in session:
        if session.get('is_admin'):
            print("   → /admin/dashboard로 리디렉션")
            return redirect('/admin/dashboard')
        else:
            print("   → /dashboard로 리디렉션")
            return redirect('/dashboard')
    
    print("   → 로그인 페이지 표시")
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    """회원가입 페이지"""
    if request.method == 'POST':
        name = request.form['name']
        team = request.form['team']
        employee_number = request.form['employee_number']
        password = request.form['password']

        if len(password) < 8:
            flash('비밀번호는 8자리 이상이어야 합니다.')
            return render_template('register.html')

        if not employee_number.startswith('N'):
            employee_number = 'N' + employee_number
            
        if len(employee_number) != 8:
            flash('사번은 7자리 숫자여야 합니다.')
            return render_template('register.html')

        try:
            int(employee_number[1:])
        except ValueError:
            flash('사번 형식이 올바르지 않습니다.')
            return render_template('register.html')

        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            
            cursor.execute('SELECT id FROM users WHERE employee_id = %s', (employee_number,))
            if cursor.fetchone():
                flash('이미 등록된 사번입니다.')
                conn.close()
                return render_template('register.html')

            hashed_password = generate_password_hash(password)
            cursor.execute('INSERT INTO users (name, employee_id, team, password) VALUES (%s, %s, %s, %s)',
                          (name, employee_number, team, hashed_password))
            
            conn.commit()
            conn.close()
            flash('회원가입이 완료되었습니다. 관리자 승인 후 이용 가능합니다.')
            return redirect('/')
            
        except Exception as e:
            flash('회원가입 중 오류가 발생했습니다.')
            return render_template('register.html')

    return render_template('register.html')

@app.route('/login', methods=['POST'])
def login():
    """로그인 처리"""
    log_session_debug('login_start')
    
    try:
        employee_id = request.form.get('employee_id', '').strip()
        password = request.form.get('password', '').strip()

        print(f"🔍 로그인 시도: '{employee_id}'")

        if not employee_id or not password:
            flash('아이디와 비밀번호를 입력해주세요.')
            return redirect('/')

        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('SELECT id, name, employee_id, password, is_approved FROM users WHERE employee_id = %s', (employee_id,))
        user = cursor.fetchone()
        
        if user and check_password_hash(user[3], password):
            print(f"✅ 비밀번호 확인 성공: {user[1]}")
            
            if user[4] == 0:
                flash('관리자 승인 대기 중입니다.')
                conn.close()
                return redirect('/')

            # 세션 설정 강화
            session.clear()
            session['user_id'] = user[0]
            session['user_name'] = user[1]
            session['employee_id'] = user[2]
            session['is_admin'] = (employee_id == 'admin')
            session.permanent = True

            conn.close()

            print("✅ 세션 설정 완료:")
            log_session_debug('login_success')

            # 로그인 후 리다이렉트
            if session['is_admin']:
                print("🎯 관리자로 로그인 - /admin/dashboard로 이동")
                return redirect('/admin/dashboard')
            else:
                print("🎯 일반 사용자로 로그인 - /dashboard로 이동")
                return redirect('/dashboard')
        else:
            print("❌ 로그인 실패")
            flash('아이디 또는 비밀번호가 잘못되었습니다.')

        conn.close()
        return redirect('/')
            
    except Exception as e:
        print(f"❌ 로그인 처리 중 오류: {str(e)}")
        flash('로그인 중 오류가 발생했습니다. 다시 시도해주세요.')
        return redirect('/')

@app.route('/admin/dashboard')
def admin_dashboard():
    """관리자 전용 대시보드 - 수정된 버전"""
    log_session_debug('/admin/dashboard')
    
    if 'user_id' not in session:
        print("   → 세션 없음, /로 리디렉션")
        flash('로그인이 필요합니다.')
        return redirect('/')

    if not session.get('is_admin'):
        print("   → 관리자 권한 없음, /dashboard로 리디렉션")
        flash('관리자 권한이 필요합니다.')
        return redirect('/dashboard')

    print("   → 관리자 대시보드 정상 표시")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 🔧 수정: SQL 쿼리 단순화
        cursor.execute("SELECT id, name, employee_id, team, is_approved, created_at FROM users WHERE employee_id != %s ORDER BY created_at DESC", ('admin',))
        users = cursor.fetchall()
        
        # 재고 통계 - 단순화
        cursor.execute(
            '''SELECT COUNT(*)
               FROM inventory
               WHERE warehouse = %s AND category = %s''',
            (DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY)
        )
        result = cursor.fetchone()
        total_items = result[0] if result else 0
        
        cursor.execute(
            '''SELECT SUM(quantity)
               FROM inventory
               WHERE warehouse = %s AND category = %s''',
            (DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY)
        )
        result = cursor.fetchone() 
        total_quantity = result[0] if result and result[0] else 0
        
        cursor.execute(
            '''SELECT COUNT(*)
               FROM inventory
               WHERE warehouse = %s AND category = %s''',
            (DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY)
        )
        diy_count_result = cursor.fetchone()
        diy_count = diy_count_result[0] if diy_count_result else 0
        
        conn.close()
        
        # 안전한 데이터 구조
        warehouse_dict = {DIY_ACTIVE_LABEL: diy_count}
        
        return render_template('admin_dashboard.html', 
                             users=users or [],
                             total_items=total_items,
                             total_quantity=total_quantity,
                             warehouse_stats=warehouse_dict)
        
    except Exception as e:
        print(f"❌ 관리자 대시보드 상세 오류: {type(e).__name__}: {str(e)}")
        # 🔧 무한 루프 방지: 간단한 HTML 반환
        return f"""
        <html>
        <head><title>관리자 대시보드</title></head>
        <body style="font-family: Arial, sans-serif; padding: 20px;">
            <h1>🔧 관리자 대시보드 (임시)</h1>
            <p>환영합니다, {session.get('user_name')}님!</p>
            <p>시스템에 일시적인 문제가 있습니다.</p>
            <p>오류: {str(e)}</p>
            <a href="/logout">로그아웃</a>
        </body>
        </html>
        """

@app.route('/dashboard')
def user_dashboard():
    """사용자 대시보드"""
    log_session_debug('/dashboard')
    
    if 'user_id' not in session:
        print("   → 세션 없음, /로 리디렉션")
        return redirect('/')

    if session.get('is_admin'):
        print("   → 관리자 감지, /admin/dashboard로 리디렉션")
        return redirect('/admin/dashboard')

    print("   → 사용자 대시보드 정상 표시")
    return render_template('user_dashboard.html', warehouses=WAREHOUSES)

@app.route('/admin/warehouse')
def admin_warehouse():
    """관리자용 창고 관리 페이지"""
    if 'user_id' not in session:
        flash('로그인이 필요합니다.')
        return redirect('/')
    
    if not session.get('is_admin'):
        flash('관리자 권한이 필요합니다.')
        return redirect('/dashboard')
    
    print("✅ 관리자 창고 관리 페이지 접근 성공")
    
    # 관리자는 모든 창고에 접근 가능
    return render_template('user_dashboard.html', warehouses=WAREHOUSES)


@app.route('/preparing')
def preparing():
    if 'user_id' not in session:
        return redirect('/')
    return render_template('preparing.html', warehouse_name=DIY_PREPARING_LABEL)

# ========
# NEW: Access 관리 관련 라우트들
# ========
@app.route('/warehouse/<warehouse_name>/access')
def access_inventory(warehouse_name):
    """Access 관리 - 기타 부품 재고 관리 페이지"""
    if 'user_id' not in session:
        return redirect('/')

    db_warehouse_name = get_db_warehouse_from_slug(warehouse_name)
    if not db_warehouse_name:
        return render_template('preparing.html', warehouse_name=DIY_PREPARING_LABEL)

    print(f"🔍 Access 관리 접근: {warehouse_name}, 사용자: {session.get('user_name')}")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('''SELECT i.id, i.category, i.part_name, i.quantity, i.last_modifier, i.last_modified,
                                COUNT(p.id) as photo_count
                         FROM inventory i
                         LEFT JOIN photos p ON i.id = p.inventory_id
                         WHERE i.warehouse = %s AND i.category = %s
                         GROUP BY i.id, i.category, i.part_name, i.quantity, i.last_modifier, i.last_modified
                         ORDER BY i.id''', (db_warehouse_name, "기타"))
        
        raw_inventory = cursor.fetchall()
        conn.close()
        
        # 🔧 날짜 형식 변환 처리 (datetime 오류 완전 해결)
        inventory = []
        for item in raw_inventory:
            item_list = list(item)
            if item_list[5]:  # last_modified가 존재하면
                if isinstance(item_list[5], str):
                    # 이미 문자열이면 그대로 사용
                    pass
                else:
                    # datetime 객체면 문자열로 변환
                    item_list[5] = item_list[5].strftime('%Y-%m-%d %H:%M:%S')
            inventory.append(item_list)
        
        print(f"✅ Access 관리 재고 데이터 조회 성공: {len(inventory)}개 항목")
        
        return render_template('access_inventory.html',
                               warehouse_name=DIY_ACTIVE_LABEL,
                               warehouse_slug=DIY_ACTIVE_SLUG,
                               warehouse_db_name=db_warehouse_name,
                               inventory=inventory,
                               is_admin=session.get('is_admin', False))
                               
    except Exception as e:
        print(f"❌ access_inventory 오류: {type(e).__name__}: {str(e)}")
        flash('재고 정보를 불러오는 중 오류가 발생했습니다.')
        
        # 🔧 관리자/사용자 구분하여 안전한 리디렉션 (무한 루프 방지)
        if session.get('is_admin'):
            return redirect('/admin/warehouse')
        else:
            return redirect('/dashboard')

# app.py의 수정된 부분들만 표시

@app.route('/save_receipt_with_details', methods=['POST'])
def save_receipt_with_details():
    """인수증 저장 (상세 정보 포함) - 수정된 버전"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '로그인이 필요합니다.'})
    
    try:
        data = request.get_json()
        receipt_date = data.get('date')
        receipt_type = data.get('type')
        warehouse_name = data.get('warehouse')
        deliverer_dept = data.get('deliverer_dept')
        deliverer_name = data.get('deliverer_name')
        receiver_dept = data.get('receiver_dept')
        receiver_name = data.get('receiver_name')
        purpose = data.get('purpose')
        items = data.get('items', [])
        
        print(f"📋 인수증 저장 시도 - 창고: {warehouse_name}, 타입: {receipt_type}, 아이템 수: {len(items)}")
        
        # 상세 정보를 포함한 데이터 구조
        detailed_data = {
            'warehouse': warehouse_name,
            'deliverer': {'dept': deliverer_dept, 'name': deliverer_name},
            'receiver': {'dept': receiver_dept, 'name': receiver_name},
            'purpose': purpose,
            'items': items
        }
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # JSON 형태로 저장 (문자열 변환 시 따옴표 처리 개선)
        items_data_json = json.dumps(detailed_data, ensure_ascii=False)
        
        cursor.execute('''
            INSERT INTO delivery_receipts 
            (receipt_date, receipt_type, items_data, created_by) 
            VALUES (%s, %s, %s, %s)
        ''', (receipt_date, receipt_type, items_data_json, session['user_name']))
        
        conn.commit()
        
        # 저장된 ID 가져오기
        cursor.execute('SELECT LASTVAL()')
        receipt_id = cursor.fetchone()[0]
        
        conn.close()
        
        print(f"✅ 인수증 저장 완료 - ID: {receipt_id}")
        
        return jsonify({
            'success': True,
            'receipt_id': receipt_id,
            'message': '인수증이 저장되었습니다.'
        })
        
    except Exception as e:
        print(f"❌ 인수증 저장 오류: {e}")
        return jsonify({'success': False, 'message': f'인수증 저장 중 오류가 발생했습니다: {str(e)}'})

# receipt_history 라우트에 추가할 코드

@app.route('/receipt_history/<warehouse_name>')
def receipt_history(warehouse_name):
    """인수증 이력 관리 페이지 - 오류 수정 버전"""
    
    print("현재 세션 키들:", list(session.keys()))
    if 'user_name' not in session and 'user_id' not in session:
        return redirect('/')
    
    db_warehouse_name = get_db_warehouse_from_slug(warehouse_name)
    if not db_warehouse_name:
        return render_template('preparing.html', warehouse_name=DIY_PREPARING_LABEL)

    print(f"🔍 인수증 이력 조회 시작 - 창고: {warehouse_name}")
    
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # ID 포함하여 조회 (삭제 기능용)
        cursor.execute('''
            SELECT id, receipt_date, receipt_type, items_data, created_by, created_at
            FROM delivery_receipts
            WHERE items_data LIKE %s
            ORDER BY receipt_date DESC, created_at DESC
            LIMIT 20
        ''', (f'%{db_warehouse_name}%',))
        
        receipts = cursor.fetchall()
        conn.close()
        
        print(f"📋 조회된 인수증: {len(receipts)}개")
        
        # 안전한 파싱 - 비고 정보 개선
        parsed_receipts = []
        
        for receipt in receipts:
            try:
                receipt_id = receipt[0]
                receipt_date = receipt[1]
                receipt_type = receipt[2]
                items_data = receipt[3]
                created_by = receipt[4]
                
                # 날짜 처리
                if hasattr(receipt_date, 'strftime'):
                    formatted_date = receipt_date.strftime('%Y-%m-%d')
                else:
                    formatted_date = str(receipt_date) if receipt_date else ''
                
                print(f"🔍 처리 중인 인수증: {receipt_id}, 날짜: {formatted_date}, 타입: {receipt_type}")
                
                # items_data 안전하게 파싱
                items_list = []
                
                if items_data:
                    try:
                        if isinstance(items_data, str):
                            parsed_data = json.loads(items_data)
                        else:
                            parsed_data = items_data
                        
                        print(f"📊 파싱된 데이터 타입: {type(parsed_data)}")
                        
                        # parsed_data가 딕셔너리이고 'items' 키가 있는 경우
                        if isinstance(parsed_data, dict) and 'items' in parsed_data:
                            items_raw = parsed_data['items']
                            print(f"📦 아이템 개수: {len(items_raw) if isinstance(items_raw, list) else 0}")
                            
                            if isinstance(items_raw, list):
                                for item in items_raw:
                                    if isinstance(item, dict):
                                        part_name = item.get('part_name', item.get('name', '알 수 없음'))
                                        quantity = item.get('quantity', item.get('qty', 0))
                                        
                                        # 비고 생성 (수정된 함수 호출)
                                        try:
                                            remark = generate_quantity_remark(warehouse_name, part_name, quantity, receipt_type)
                                        except Exception as remark_error:
                                            print(f"비고 생성 실패: {remark_error}")
                                            remark = f"{receipt_type} {quantity}개"
                                        
                                        items_list.append({
                                            'part_name': part_name,
                                            'quantity': quantity,
                                            'deliverer_dept': item.get('deliverer_dept', '-'),
                                            'deliverer_name': item.get('deliverer_name', '-'),
                                            'receiver_dept': item.get('receiver_dept', '-'),
                                            'receiver_name': item.get('receiver_name', '-'),
                                            'purpose': item.get('purpose', '-'),
                                            'remark': remark
                                        })
                                    else:
                                        items_list.append({
                                            'part_name': str(item),
                                            'quantity': 0,
                                            'deliverer_dept': '-',
                                            'deliverer_name': '-',
                                            'receiver_dept': '-',
                                            'receiver_name': '-',
                                            'purpose': '-',
                                            'remark': '-'
                                        })
                        
                        # parsed_data가 리스트인 경우 (구 형식)
                        elif isinstance(parsed_data, list):
                            print("📦 구 형식 리스트 데이터 처리")
                            for item in parsed_data:
                                if isinstance(item, dict):
                                    part_name = item.get('part_name', item.get('name', '알 수 없음'))
                                    quantity = item.get('quantity', item.get('qty', 0))
                                    
                                    try:
                                        remark = generate_quantity_remark(warehouse_name, part_name, quantity, receipt_type)
                                    except Exception as remark_error:
                                        print(f"비고 생성 실패: {remark_error}")
                                        remark = f"{receipt_type} {quantity}개"
                                    
                                    items_list.append({
                                        'part_name': part_name,
                                        'quantity': quantity,
                                        'deliverer_dept': item.get('deliverer_dept', '-'),
                                        'deliverer_name': item.get('deliverer_name', '-'),
                                        'receiver_dept': item.get('receiver_dept', '-'),
                                        'receiver_name': item.get('receiver_name', '-'),
                                        'purpose': item.get('purpose', '-'),
                                        'remark': remark
                                    })
                                else:
                                    items_list.append({
                                        'part_name': str(item),
                                        'quantity': 0,
                                        'deliverer_dept': '-',
                                        'deliverer_name': '-',
                                        'receiver_dept': '-',
                                        'receiver_name': '-',
                                        'purpose': '-',
                                        'remark': '-'
                                    })
                        else:
                            print(f"⚠️ 알 수 없는 데이터 형식: {type(parsed_data)}")
                            items_list = [{
                                'part_name': '알 수 없는 형식',
                                'quantity': 0,
                                'deliverer_dept': '-',
                                'deliverer_name': '-',
                                'receiver_dept': '-',
                                'receiver_name': '-',
                                'purpose': '-',
                                'remark': '데이터 형식 오류'
                            }]
                        
                    except (json.JSONDecodeError, TypeError, AttributeError) as e:
                        print(f"⚠️ items_data JSON 파싱 오류: {e}")
                        items_list = [{
                            'part_name': 'JSON 파싱 오류',
                            'quantity': 0,
                            'deliverer_dept': '-',
                            'deliverer_name': '-',
                            'receiver_dept': '-',
                            'receiver_name': '-',
                            'purpose': '-',
                            'remark': 'JSON 오류'
                        }]
                else:
                    print("⚠️ items_data가 비어있음")
                    items_list = [{
                        'part_name': '데이터 없음',
                        'quantity': 0,
                        'deliverer_dept': '-',
                        'deliverer_name': '-',
                        'receiver_dept': '-',
                        'receiver_name': '-',
                        'purpose': '-',
                        'remark': '데이터 없음'
                    }]
                
                receipt_dict = {
                    'id': receipt_id,
                    'date': formatted_date,
                    'type': receipt_type or 'unknown',
                    'receipt_items': items_list,
                    'created_by': created_by or '미설정'
                }
                
                parsed_receipts.append(receipt_dict)
                print(f"✅ 인수증 {receipt_id} 파싱 완료: {len(items_list)}개 아이템")
                
            except Exception as e:
                print(f"⚠️ 인수증 전체 파싱 오류: {e}")
                import traceback
                print(f"상세 오류: {traceback.format_exc()}")
                
                parsed_receipts.append({
                    'id': receipt[0] if len(receipt) > 0 else 0,
                    'date': '날짜 오류',
                    'type': 'unknown',
                    'receipt_items': [{
                        'part_name': '전체 오류 발생',
                        'quantity': 0,
                        'deliverer_dept': '-',
                        'deliverer_name': '-',
                        'receiver_dept': '-',
                        'receiver_name': '-',
                        'purpose': '-',
                        'remark': '전체 오류'
                    }],
                    'created_by': '미설정'
                })
                continue
        
        print(f"✅ 전체 파싱 완료: {len(parsed_receipts)}개")
        
        template_vars = {
            'warehouse_name': DIY_ACTIVE_LABEL,
            'warehouse_slug': DIY_ACTIVE_SLUG,
            'receipts': parsed_receipts,
            'current_page': 1,
            'total_pages': 1,
            'total_count': len(parsed_receipts),
            'is_admin': session.get('is_admin', False)
        }
        
        return render_template('receipt_history.html', **template_vars)
        
    except Exception as e:
        print(f"❌ 인수증 이력 조회 전체 오류: {e}")
        import traceback
        print(f"상세 오류: {traceback.format_exc()}")
        flash('인수증 이력을 불러오는 중 오류가 발생했습니다.')
        return redirect(f'/warehouse/{DIY_ACTIVE_SLUG}/access')
        
def generate_quantity_remark(warehouse_name, part_name, quantity, receipt_type):
    """수량 변화 비고 생성 함수 - 올바른 버전"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 현재 재고량 조회
        cursor.execute('''
            SELECT quantity FROM inventory 
            WHERE warehouse = %s AND part_name = %s AND category = %s
        ''', (warehouse_name, part_name, "기타"))
        
        result = cursor.fetchone()
        current_qty = result[0] if result else 0
        
        conn.close()
        
        if receipt_type == 'in':
            # 입고: 현재 수량에서 입고량을 뺀 것이 입고 전 수량
            before_qty = max(0, current_qty - quantity)
            after_qty = current_qty
            return f"입고전 {before_qty}개 → 입고후 {after_qty}개"
        else:
            # 출고: 현재 수량에 출고량을 더한 것이 출고 전 수량
            before_qty = current_qty + quantity
            after_qty = current_qty
            return f"출고전 {before_qty}개 → 출고후 {after_qty}개"
            
    except Exception as e:
        print(f"비고 생성 오류: {e}")
        if receipt_type == 'in':
            return f"입고 {quantity}개"
        else:
            return f"출고 {quantity}개"
        

# 디버깅용 라우트 추가
@app.route('/debug_receipts/<warehouse_name>')
def debug_receipts(warehouse_name):
    """인수증 디버깅 페이지 (관리자 전용)"""
    if 'user_id' not in session or not session.get('is_admin'):
        flash('관리자 권한이 필요합니다.')
        return redirect('/')
    
    db_warehouse_name = get_db_warehouse_from_slug(warehouse_name)
    if not db_warehouse_name:
        return render_template('preparing.html', warehouse_name=DIY_PREPARING_LABEL)

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 모든 인수증 조회
        cursor.execute('SELECT id, receipt_date, receipt_type, items_data, created_by, created_at FROM delivery_receipts ORDER BY created_at DESC LIMIT 20')
        all_receipts = cursor.fetchall()
        
        # 특정 창고 인수증 조회
        cursor.execute('''
            SELECT id, receipt_date, receipt_type, items_data, created_by, created_at 
            FROM delivery_receipts 
            WHERE items_data::text LIKE %s 
            ORDER BY created_at DESC LIMIT 20
        ''', (f'%"warehouse": "{db_warehouse_name}"%',))
        warehouse_receipts = cursor.fetchall()
        
        conn.close()
        
        debug_info = {
            'warehouse_name': DIY_ACTIVE_LABEL,
            'total_receipts': len(all_receipts),
            'warehouse_receipts': len(warehouse_receipts),
            'all_receipts': all_receipts,
            'filtered_receipts': warehouse_receipts
        }
        
        return f"""
        <html>
        <head><title>인수증 디버깅 - {DIY_ACTIVE_LABEL}</title></head>
        <body style="font-family: Arial, sans-serif; padding: 20px;">
            <h1>인수증 디버깅 정보</h1>
            <h2>DIY 위치: {DIY_ACTIVE_LABEL}</h2>
            
            <h3>📊 통계</h3>
            <ul>
                <li>전체 인수증 개수: {debug_info['total_receipts']}</li>
                <li>{DIY_ACTIVE_LABEL} 인수증: {debug_info['warehouse_receipts']}</li>
            </ul>
            
            <h3>🔍 최근 {DIY_ACTIVE_LABEL} 인수증들</h3>
            <table border="1" style="border-collapse: collapse; width: 100%;">
                <tr>
                    <th>ID</th>
                    <th>날짜</th>
                    <th>타입</th>
                    <th>생성자</th>
                    <th>생성시간</th>
                    <th>데이터 미리보기</th>
                </tr>
                {''.join([f"<tr><td>{r[0]}</td><td>{r[1]}</td><td>{r[2]}</td><td>{r[4]}</td><td>{r[5]}</td><td>{str(r[3])[:100]}...</td></tr>" for r in warehouse_receipts])}
            </table>
            
            <h3>🗂️ 전체 인수증들 (최근 50개)</h3>
            <table border="1" style="border-collapse: collapse; width: 100%;">
                <tr>
                    <th>ID</th>
                    <th>날짜</th>
                    <th>타입</th>
                    <th>생성자</th>
                    <th>생성시간</th>
                    <th>데이터 미리보기</th>
                </tr>
                {''.join([f"<tr><td>{r[0]}</td><td>{r[1]}</td><td>{r[2]}</td><td>{r[4]}</td><td>{r[5]}</td><td>{str(r[3])[:100]}...</td></tr>" for r in all_receipts])}
            </table>
            
            <br><br>
            <a href="/warehouse/{DIY_ACTIVE_SLUG}/access">← 재고 관리로 돌아가기</a>
        </body>
        </html>
        """
        
    except Exception as e:
        return f"디버깅 오류: {str(e)}"
@app.route('/add_access_inventory_item', methods=['POST'])
def add_access_inventory_item():
    """Access 관리 - 재고 아이템 추가 (관리자 전용)"""
    if 'user_id' not in session or not session.get('is_admin'):
        flash('관리자 권한이 필요합니다.')
        return redirect('/')

    warehouse_name = request.form['warehouse_name']
    category = request.form['category']
    part_name = request.form['part_name']
    quantity = int(request.form['quantity'])
    korea_time = get_korea_time().strftime('%Y-%m-%d %H:%M:%S')

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('INSERT INTO inventory (warehouse, category, part_name, quantity, last_modifier, last_modified) VALUES (%s, %s, %s, %s, %s, %s)',
                      (warehouse_name, category, part_name, quantity, session['user_name'], korea_time))
        
        conn.commit()
        conn.close()
        flash('재고 아이템이 추가되었습니다.')
        
    except Exception as e:
        flash('재고 추가 중 오류가 발생했습니다.')
    
    return redirect(f'/warehouse/{get_slug_from_db_warehouse(warehouse_name)}/access')

@app.route('/delivery_receipt/<warehouse_name>')
def delivery_receipt_form(warehouse_name):
    """인수증 생성 페이지"""
    if 'user_id' not in session:
        return redirect('/')
    
    db_warehouse_name = get_db_warehouse_from_slug(warehouse_name)
    if not db_warehouse_name:
        return render_template('preparing.html', warehouse_name=DIY_PREPARING_LABEL)
    return render_template('delivery_receipt.html', warehouse_name=DIY_ACTIVE_LABEL, warehouse_slug=DIY_ACTIVE_SLUG, warehouse_db_name=db_warehouse_name)

@app.route('/get_inventory_changes', methods=['POST'])
def get_inventory_changes():
    """특정 날짜의 입고/출고 내역 조회"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '로그인이 필요합니다.'})
    
    try:
        data = request.get_json()
        target_date = data.get('date')
        change_type = data.get('type')  # 'in' 또는 'out'
        warehouse_name = data.get('warehouse')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 해당 날짜의 변경 내역 조회
        cursor.execute('''
            SELECT h.inventory_id, i.part_name, h.quantity_change, h.modifier_name, h.modified_at
            FROM inventory_history h
            JOIN inventory i ON h.inventory_id = i.id
            WHERE DATE(h.modified_at AT TIME ZONE 'Asia/Seoul') = %s
            AND h.change_type = %s
            AND i.warehouse = %s
            AND i.category = %s
            ORDER BY h.modified_at DESC
        ''', (target_date, change_type, warehouse_name, "기타"))
        
        changes = cursor.fetchall()
        conn.close()
        
        # 데이터 포맷팅
        formatted_changes = []
        for change in changes:
            formatted_changes.append({
                'inventory_id': change[0],
                'part_name': change[1],
                'quantity': abs(change[2]),  # 절댓값으로 표시
                'modifier': change[3],
                'time': change[4].strftime('%H:%M') if change[4] else ''
            })
        
        return jsonify({
            'success': True,
            'changes': formatted_changes
        })
        
    except Exception as e:
        print(f"❌ 재고 변경 내역 조회 오류: {e}")
        return jsonify({'success': False, 'message': '데이터 조회 중 오류가 발생했습니다.'})

@app.route('/save_delivery_receipt', methods=['POST'])
def save_delivery_receipt():
    """인수증 저장"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '로그인이 필요합니다.'})
    
    try:
        data = request.get_json()
        receipt_date = data.get('date')
        receipt_type = data.get('type')
        items_data = data.get('items', [])
        signature_data = data.get('signature')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 인수증 데이터 저장
        cursor.execute('''
            INSERT INTO delivery_receipts 
            (receipt_date, receipt_type, items_data, signature_data, created_by) 
            VALUES (%s, %s, %s, %s, %s)
        ''', (receipt_date, receipt_type, str(items_data), signature_data, session['user_name']))
        
        conn.commit()
        receipt_id = cursor.lastrowid if cursor.lastrowid else cursor.fetchone()[0]
        conn.close()
        
        return jsonify({
            'success': True,
            'receipt_id': receipt_id,
            'message': '인수증이 저장되었습니다.'
        })
        
    except Exception as e:
        print(f"❌ 인수증 저장 오류: {e}")
        return jsonify({'success': False, 'message': '인수증 저장 중 오류가 발생했습니다.'})

@app.route('/send_delivery_receipt', methods=['POST'])
def send_delivery_receipt():
    """인수증 이메일 발송"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '로그인이 필요합니다.'})
    
    try:
        data = request.get_json()
        to_emails = data.get('emails', [])
        receipt_data = data.get('receipt_data', {})
        
        if not to_emails:
            return jsonify({'success': False, 'message': '수신자 이메일을 입력해주세요.'})
        
        # 이메일 HTML 생성
        receipt_type_korean = "입고" if receipt_data.get('type') == 'in' else "출고"
        
        html_content = f"""
        <!DOCTYPE html>
        <html lang="ko">
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                .header {{ text-align: center; margin-bottom: 30px; }}
                .receipt-info {{ background: #f8f9fa; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
                .items-table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
                .items-table th, .items-table td {{ border: 1px solid #ddd; padding: 10px; text-align: left; }}
                .items-table th {{ background-color: #f2f2f2; }}
                .signature {{ text-align: center; margin-top: 30px; }}
                .signature img {{ max-width: 300px; border: 1px solid #ddd; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h2>SK오앤에스 DIY System</h2>
                <h3>{receipt_type_korean} 인수증</h3>
            </div>
            
            <div class="receipt-info">
                <p><strong>일자:</strong> {receipt_data.get('date', '')}</p>
                <p><strong>창고:</strong> {receipt_data.get('warehouse', '')}</p>
                <p><strong>구분:</strong> {receipt_type_korean}</p>
                <p><strong>작성자:</strong> {session.get('user_name', '')}</p>
            </div>
            
            <table class="items-table">
                <thead>
                    <tr>
                        <th>번호</th>
                        <th>부품명</th>
                        <th>수량</th>
                        <th>담당자</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        for i, item in enumerate(receipt_data.get('items', []), 1):
            html_content += f"""
                    <tr>
                        <td>{i}</td>
                        <td>{item.get('part_name', '')}</td>
                        <td>{item.get('quantity', '')}개</td>
                        <td>{item.get('modifier', '')}</td>
                    </tr>
            """
        
        html_content += """
                </tbody>
            </table>
        """
        
        # 전자서명이 있으면 추가
        if receipt_data.get('signature'):
            html_content += f"""
            <div class="signature">
                <p><strong>전자서명:</strong></p>
                <img src="{receipt_data.get('signature')}" alt="전자서명">
            </div>
            """
        
        html_content += """
            <p style="text-align: center; margin-top: 30px; color: #666; font-size: 12px;">
                본 인수증은 SK오앤에스 DIY System에서 자동으로 생성되었습니다.
            </p>
        </body>
        </html>
        """
        
        # 이메일 발송
        subject = f"[SK오앤에스] {receipt_type_korean} 인수증 - {receipt_data.get('date', '')}"
        success, message = send_email(to_emails, subject, html_content)
        
        return jsonify({
            'success': success,
            'message': message
        })
        
    except Exception as e:
        print(f"❌ 인수증 이메일 발송 오류: {e}")
        return jsonify({'success': False, 'message': f'이메일 발송 중 오류가 발생했습니다: {str(e)}'})

# ========
# 기존 라우트들 계속 (변경사항 없음)
# ========
@app.route('/approve_user/<int:user_id>')
def approve_user(user_id):
    """사용자 승인 (관리자 전용)"""
    if 'user_id' not in session or not session.get('is_admin'):
        flash('관리자 권한이 필요합니다.')
        return redirect('/')

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('UPDATE users SET is_approved = %s WHERE id = %s', (1, user_id))
        conn.commit()
        conn.close()
        flash('사용자가 승인되었습니다.')
        
    except Exception as e:
        flash('사용자 승인 중 오류가 발생했습니다.')
    
    return redirect('/admin/dashboard')

@app.route('/delete_user/<int:user_id>')
def delete_user(user_id):
    """사용자 삭제 (관리자 전용)"""
    if 'user_id' not in session or not session.get('is_admin'):
        flash('관리자 권한이 필요합니다.')
        return redirect('/')

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('SELECT name, employee_id FROM users WHERE id = %s AND employee_id != %s', (user_id, 'admin'))
        user = cursor.fetchone()
        
        if user:
            cursor.execute('DELETE FROM users WHERE id = %s', (user_id,))
            conn.commit()
            flash(f'사용자 {user[0]}({user[1]})가 삭제되었습니다.')
        else:
            flash('삭제할 수 없는 사용자입니다.')
        
        conn.close()
        
    except Exception as e:
        flash('사용자 삭제 중 오류가 발생했습니다.')
    
    return redirect('/admin/dashboard')

@app.route('/warehouse/<warehouse_name>')
def warehouse(warehouse_name):
    """창고 선택 페이지"""
    if 'user_id' not in session:
        return redirect('/')

    db_warehouse_name = get_db_warehouse_from_slug(warehouse_name)
    if not db_warehouse_name:
        return render_template('preparing.html', warehouse_name=DIY_PREPARING_LABEL)

    return render_template(
        'warehouse.html',
        warehouse_name=get_display_warehouse_from_slug(warehouse_name),
        warehouse_slug=warehouse_name,
        diy_intro_text=get_diy_intro_text(warehouse_name)
    )

@app.route('/warehouse/<warehouse_name>/electric')
def electric_inventory(warehouse_name):
    """DIY 점검 대상 목록 페이지"""
    if 'user_id' not in session:
        return redirect('/')

    db_warehouse_name = get_db_warehouse_from_slug(warehouse_name)
    if not db_warehouse_name:
        return render_template('preparing.html', warehouse_name=DIY_PREPARING_LABEL)

    warehouse_label = get_display_warehouse_from_slug(warehouse_name)
    use_equipment_units = uses_equipment_units(warehouse_name)
    allowed_site_names = FACILITY_INSPECTION_SITE_NAMES if warehouse_name == DIY_FACILITY_SLUG else []

    search_query = request.args.get('q', '').strip()
    status_filter = request.args.get('status', '').strip()
    if status_filter not in {'completed', 'pending'}:
        status_filter = ''

    print(f"🔍 DIY 작업 관리 접근: {warehouse_name}, 사용자: {session.get('user_name')}")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        base_where_conditions = ["i.warehouse = %s", "i.category = %s"]
        base_params = [db_warehouse_name, DIY_CHECKLIST_CATEGORY]
        if allowed_site_names:
            base_where_conditions.append("(i.part_name = %s OR i.part_name = %s)")
            base_params.extend(allowed_site_names[:2])
        if search_query:
            base_where_conditions.append("(i.part_name ILIKE %s OR COALESCE(r.inspector_name, '') ILIKE %s)")
            base_params.append(f'%{search_query}%')
            base_params.append(f'%{search_query}%')

        count_query = f'''
            SELECT
                SUM(CASE WHEN r.inspected_at IS NOT NULL THEN 1 ELSE 0 END) AS completed_count,
                SUM(CASE WHEN r.inspected_at IS NULL THEN 1 ELSE 0 END) AS pending_count
            FROM inventory i
            LEFT JOIN (
                SELECT DISTINCT ON (inventory_id)
                       id,
                       inventory_id,
                       inspector_name,
                       inspected_at
                FROM inspection_records
                ORDER BY inventory_id, inspected_at DESC
            ) r ON i.id = r.inventory_id
            WHERE {' AND '.join(base_where_conditions)}
        '''
        cursor.execute(count_query, base_params)
        count_row = cursor.fetchone() or (0, 0)
        completed_count = int(count_row[0] or 0)
        pending_count = int(count_row[1] or 0)

        where_conditions = list(base_where_conditions)
        params = list(base_params)
        if status_filter == 'completed':
            where_conditions.append("r.inspected_at IS NOT NULL")
        elif status_filter == 'pending':
            where_conditions.append("r.inspected_at IS NULL")

        cursor.execute(f'''
            SELECT i.id,
                   i.part_name,
                   r.id,
                   r.inspector_name,
                   r.inspected_at
            FROM inventory i
            LEFT JOIN (
                SELECT DISTINCT ON (inventory_id)
                       id,
                       inventory_id,
                       inspector_name,
                       inspected_at
                FROM inspection_records
                ORDER BY inventory_id, inspected_at DESC
            ) r ON i.id = r.inventory_id
            WHERE {' AND '.join(where_conditions)}
            ORDER BY i.id
        ''', params)
        
        raw_inventory = cursor.fetchall()
        conn.close()
        
        checklist_targets = []
        for row in raw_inventory:
            target_id = row[0]
            site_name = row[1] or '미입력'
            latest_record_id = row[2]
            inspector_name = row[3] or ''
            inspected_at = row[4]
            if inspected_at and not isinstance(inspected_at, str):
                inspected_at = inspected_at.strftime('%Y-%m-%d %H:%M:%S')
            checklist_targets.append({
                'id': target_id,
                'site_name': site_name,
                'inspector_name': inspector_name,
                'inspected_at': inspected_at or '',
                'status': '작업 완료' if inspected_at else '작업 미완료',
                'is_completed': bool(inspected_at),
                'latest_record_id': latest_record_id,
                'detail_url': (
                    url_for('inspection_units', warehouse_name=warehouse_name, item_id=target_id)
                    if use_equipment_units
                    else url_for('inspection_detail', warehouse_name=warehouse_name, item_id=target_id)
                ),
                'hangul_url': (
                    url_for('export_facility_report_docx', warehouse_name=warehouse_name, item_id=target_id)
                    if (warehouse_name == DIY_FACILITY_SLUG and inspected_at)
                    else ''
                )
            })
        
        print(f"✅ 점검 대상 조회 성공: {len(checklist_targets)}개 항목")
        
        return render_template('electric_inventory.html',
                               warehouse_name=warehouse_label,
                               warehouse_slug=warehouse_name,
                               warehouse_db_name=db_warehouse_name,
                               checklist_targets=checklist_targets,
                               search_query=search_query,
                               status_filter=status_filter,
                               completed_count=completed_count,
                               pending_count=pending_count,
                               inspection_items=INSPECTION_ITEMS,
                               is_admin=session.get('is_admin', False),
                               show_admin_sites_button=(session.get('is_admin', False) and use_equipment_units),
                               show_excel_export=use_equipment_units)
                               
    except Exception as e:
        print(f"❌ electric_inventory 오류: {type(e).__name__}: {str(e)}")
        flash('재고 정보를 불러오는 중 오류가 발생했습니다.')
        
        # 🔧 관리자/사용자 구분하여 안전한 리디렉션 (무한 루프 방지)
        if session.get('is_admin'):
            return redirect('/admin/warehouse')
        else:
            return redirect('/dashboard')


@app.route('/warehouse/<warehouse_name>/inspection/<int:item_id>', methods=['GET', 'POST'])
def inspection_detail(warehouse_name, item_id):
    """점검 상세 페이지(보기/수정) 및 저장"""
    if 'user_id' not in session:
        return redirect('/')

    db_warehouse_name = get_db_warehouse_from_slug(warehouse_name)
    if not db_warehouse_name:
        return render_template('preparing.html', warehouse_name=DIY_PREPARING_LABEL)
    facility_mode = is_facility_suitability(warehouse_name)

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            '''SELECT id, part_name
               FROM inventory
               WHERE id = %s AND warehouse = %s AND category = %s''',
            (item_id, db_warehouse_name, DIY_CHECKLIST_CATEGORY)
        )
        item = cursor.fetchone()
        if not item:
            conn.close()
            flash('점검 대상을 찾을 수 없습니다.')
            return redirect(f'/warehouse/{warehouse_name}/electric')

        cursor.execute(
            '''SELECT id, site_name, inspector_name, inspected_at, checklist_data, memo
               FROM inspection_records
               WHERE inventory_id = %s
               ORDER BY inspected_at DESC
               LIMIT 1''',
            (item_id,)
        )
        latest_record = cursor.fetchone()

        latest_checklist_by_no = {}
        latest_photos = {}
        latest_facility_payload = build_default_facility_payload()
        latest_facility_capture_photos = {}
        if latest_record:
            if facility_mode:
                latest_facility_payload = parse_facility_payload(latest_record[4])
                latest_facility_capture_photos = fetch_facility_capture_photos(cursor, latest_record[0])
            else:
                if latest_record[4]:
                    try:
                        checklist_raw = latest_record[4]
                        checklist_data = checklist_raw if isinstance(checklist_raw, list) else json.loads(checklist_raw)
                        for row in checklist_data:
                            checkpoint_no = int(row.get('checkpoint_no', 0))
                            latest_checklist_by_no[checkpoint_no] = row.get('result', 'ok')
                    except Exception:
                        latest_checklist_by_no = {}

                cursor.execute(
                    '''SELECT checkpoint_no, checkpoint_name, phase, filename, file_size, supabase_url
                       FROM inspection_photos
                       WHERE record_id = %s''',
                    (latest_record[0],)
                )
                for checkpoint_no, checkpoint_name, phase, filename, file_size, supabase_url in cursor.fetchall():
                    if checkpoint_no not in latest_photos:
                        latest_photos[checkpoint_no] = {}
                    latest_photos[checkpoint_no][phase] = {
                        'checkpoint_name': checkpoint_name,
                        'filename': filename,
                        'file_size': file_size,
                        'supabase_url': supabase_url
                    }

        if request.method == 'POST':
            korea_time = get_korea_time().strftime('%Y-%m-%d %H:%M:%S')
            inspector_name = session.get('user_name', '미설정')
            requested_site_name = request.form.get('site_name', '').strip()
            if session.get('is_admin'):
                site_name = requested_site_name or (item[1] or '미입력')
            else:
                site_name = (latest_record[1] if latest_record and latest_record[1] else item[1] or '미입력')
            plain_memo = request.form.get('memo', '').strip()
            edit_mode = request.form.get('edit_mode') == '1'
            record_id = request.form.get('record_id', '').strip()
            update_existing = (
                edit_mode and
                record_id.isdigit() and
                latest_record and
                int(record_id) == latest_record[0]
            )
            if facility_mode:
                facility_payload = build_facility_payload_from_form(request.form)
                checklist_json = json.dumps(facility_payload, ensure_ascii=False)
                existing_capture_map = latest_facility_capture_photos if update_existing else {}

                if update_existing:
                    record_id_int = int(record_id)
                    cursor.execute(
                        '''UPDATE inspection_records
                           SET site_name = %s,
                               inspector_name = %s,
                               inspected_at = %s,
                               status = %s,
                               checklist_data = %s,
                               memo = %s
                           WHERE id = %s AND inventory_id = %s''',
                        (site_name, inspector_name, korea_time, '작업 완료', checklist_json, plain_memo, record_id_int, item_id)
                    )
                    target_record_id = record_id_int
                else:
                    cursor.execute(
                        '''INSERT INTO inspection_records
                           (inventory_id, warehouse, site_name, inspector_name, inspected_at, status, checklist_data, memo)
                           VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                           RETURNING id''',
                        (item_id, db_warehouse_name, site_name, inspector_name, korea_time, '작업 완료', checklist_json, plain_memo)
                    )
                    target_record_id = cursor.fetchone()[0]

                cursor.execute('DELETE FROM facility_capture_photos WHERE record_id = %s', (target_record_id,))

                for slot in FACILITY_CAPTURE_SLOTS:
                    slot_no = slot['slot_no']
                    photo_file = request.files.get(f'facility_capture_{slot_no}')
                    saved_photo = None
                    if photo_file and photo_file.filename:
                        saved_photo = save_facility_capture_photo(photo_file, item_id, slot_no)
                    elif existing_capture_map.get(slot_no):
                        saved_photo = {
                            'filename': existing_capture_map[slot_no]['filename'],
                            'file_size': existing_capture_map[slot_no]['file_size'],
                            'supabase_url': existing_capture_map[slot_no]['supabase_url']
                        }

                    if saved_photo:
                        cursor.execute(
                            '''INSERT INTO facility_capture_photos
                               (record_id, slot_no, point_title, detail_text, filename, file_size, supabase_url)
                               VALUES (%s, %s, %s, %s, %s, %s, %s)''',
                            (
                                target_record_id,
                                slot_no,
                                slot['title'],
                                slot['detail'],
                                saved_photo['filename'],
                                saved_photo['file_size'],
                                saved_photo['supabase_url']
                            )
                        )

                cursor.execute(
                    'UPDATE inventory SET part_name = %s, last_modifier = %s, last_modified = %s WHERE id = %s',
                    (site_name, inspector_name, korea_time, item_id)
                )
                cursor.execute(
                    '''INSERT INTO inventory_history
                       (inventory_id, change_type, quantity_change, modifier_name, modified_at)
                       VALUES (%s, %s, %s, %s, %s)''',
                    (item_id, 'inspection', 0, inspector_name, korea_time)
                )

                conn.commit()
                conn.close()
                flash('점검 내용이 저장되었습니다.')
                return redirect(f'/warehouse/{warehouse_name}/electric')

            checklist_data = []
            photo_rows = []

            for checkpoint_no, checkpoint_name in INSPECTION_ITEMS:
                result_value = request.form.get(f'result_{checkpoint_no}', 'ok')
                if result_value not in {'ok', 'need', 'na'}:
                    result_value = 'ok'
                before_file = request.files.get(f'before_{checkpoint_no}')
                after_file = request.files.get(f'after_{checkpoint_no}')
                existing_before = latest_photos.get(checkpoint_no, {}).get('before') if update_existing else None
                existing_after = latest_photos.get(checkpoint_no, {}).get('after') if update_existing else None

                before_name = None
                before_size = None
                before_url = None
                if before_file and before_file.filename:
                    before_name, before_size, before_url = save_inspection_photo(before_file, item_id, checkpoint_no, 'before')
                elif existing_before:
                    before_name = existing_before['filename']
                    before_size = existing_before['file_size']
                    before_url = existing_before['supabase_url']

                after_name = None
                after_size = None
                after_url = None
                if after_file and after_file.filename:
                    after_name, after_size, after_url = save_inspection_photo(after_file, item_id, checkpoint_no, 'after')
                elif existing_after:
                    after_name = existing_after['filename']
                    after_size = existing_after['file_size']
                    after_url = existing_after['supabase_url']

                checklist_data.append({
                    'checkpoint_no': checkpoint_no,
                    'checkpoint_name': checkpoint_name,
                    'result': result_value
                })

                if before_url:
                    photo_rows.append((checkpoint_no, checkpoint_name, 'before', before_name, before_size, before_url))
                if after_url:
                    photo_rows.append((checkpoint_no, checkpoint_name, 'after', after_name, after_size, after_url))

            if update_existing:
                record_id_int = int(record_id)
                cursor.execute(
                    '''UPDATE inspection_records
                       SET site_name = %s,
                           inspector_name = %s,
                           inspected_at = %s,
                           status = %s,
                           checklist_data = %s,
                           memo = %s
                       WHERE id = %s AND inventory_id = %s''',
                    (site_name, inspector_name, korea_time, '작업 완료',
                     json.dumps(checklist_data, ensure_ascii=False), plain_memo, record_id_int, item_id)
                )
                cursor.execute('DELETE FROM inspection_photos WHERE record_id = %s', (record_id_int,))
                target_record_id = record_id_int
            else:
                cursor.execute(
                    '''INSERT INTO inspection_records
                       (inventory_id, warehouse, site_name, inspector_name, inspected_at, status, checklist_data, memo)
                       VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                       RETURNING id''',
                    (item_id, db_warehouse_name, site_name, inspector_name, korea_time, '작업 완료',
                     json.dumps(checklist_data, ensure_ascii=False), plain_memo)
                )
                target_record_id = cursor.fetchone()[0]

            for checkpoint_no, checkpoint_name, phase, filename, file_size, supabase_url in photo_rows:
                cursor.execute(
                    '''INSERT INTO inspection_photos
                       (record_id, checkpoint_no, checkpoint_name, phase, filename, file_size, supabase_url)
                       VALUES (%s, %s, %s, %s, %s, %s, %s)''',
                    (target_record_id, checkpoint_no, checkpoint_name, phase, filename, file_size, supabase_url)
                )

            cursor.execute(
                'UPDATE inventory SET part_name = %s, last_modifier = %s, last_modified = %s WHERE id = %s',
                (site_name, inspector_name, korea_time, item_id)
            )
            cursor.execute(
                '''INSERT INTO inventory_history
                   (inventory_id, change_type, quantity_change, modifier_name, modified_at)
                   VALUES (%s, %s, %s, %s, %s)''',
                (item_id, 'inspection', 0, inspector_name, korea_time)
            )

            conn.commit()
            conn.close()
            flash('점검 내용이 저장되었습니다.')
            return redirect(f'/warehouse/{warehouse_name}/electric')

        editable = (request.args.get('mode') == 'edit') or (latest_record is None)
        latest_record_dict = None
        memo_plain_text = ''
        inspected_at_str = ''
        if latest_record:
            inspected_at_value = latest_record[3]
            if inspected_at_value and not isinstance(inspected_at_value, str):
                inspected_at_str = inspected_at_value.strftime('%Y-%m-%d %H:%M:%S')
            else:
                inspected_at_str = inspected_at_value or ''
            memo_plain_text = latest_record[5] or ''

            latest_record_dict = {
                'id': latest_record[0],
                'site_name': latest_record[1] or item[1] or '미입력',
                'inspector_name': latest_record[2] or '',
                'inspected_at': inspected_at_str,
                'memo': memo_plain_text
            }

        if facility_mode and not latest_facility_payload.get('inspection_date'):
            latest_facility_payload['inspection_date'] = get_korea_time().strftime('%Y-%m-%d')

        conn.close()
        return render_template(
            'inspection_detail.html',
            warehouse_name=get_display_warehouse_from_slug(warehouse_name),
            warehouse_slug=warehouse_name,
            item_id=item[0],
            site_name=(latest_record_dict['site_name'] if latest_record_dict else (item[1] or '미입력')),
            equipment_no='',
            equipment_selector_url='',
            back_url=url_for('electric_inventory', warehouse_name=warehouse_name),
            back_button_label='목록으로',
            show_management_actions=False,
            edit_url=url_for('inspection_detail', warehouse_name=warehouse_name, item_id=item_id, mode='edit'),
            inspector_name=session.get('user_name', '미설정'),
            inspection_items=INSPECTION_ITEMS,
            editable=editable,
            latest_record=latest_record_dict,
            latest_checklist_by_no=latest_checklist_by_no,
            latest_photos=latest_photos,
            is_admin=session.get('is_admin', False),
            facility_mode=facility_mode,
            facility_sections=FACILITY_SURVEY_SECTIONS,
            facility_payload=latest_facility_payload,
            facility_capture_slots=FACILITY_CAPTURE_SLOTS,
            facility_capture_photos=latest_facility_capture_photos
        )
    except Exception as e:
        try:
            conn.rollback()
            conn.close()
        except Exception:
            pass
        print(f"❌ inspection_detail 오류: {type(e).__name__}: {str(e)}")
        flash(f'점검 페이지 처리 중 오류가 발생했습니다: {str(e)}')
        return redirect(f'/warehouse/{warehouse_name}/electric')


@app.route('/warehouse/<warehouse_name>/inspection-units/<int:item_id>', methods=['GET', 'POST'])
def inspection_units(warehouse_name, item_id):
    if 'user_id' not in session:
        return redirect('/')

    db_warehouse_name = get_db_warehouse_from_slug(warehouse_name)
    if not db_warehouse_name:
        return render_template('preparing.html', warehouse_name=DIY_PREPARING_LABEL)

    manage_mode = request.args.get('mode') == 'edit'

    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            '''SELECT id, part_name
               FROM inventory
               WHERE id = %s AND warehouse = %s AND category = %s''',
            (item_id, db_warehouse_name, DIY_CHECKLIST_CATEGORY)
        )
        item = cursor.fetchone()
        if not item:
            conn.close()
            flash('점검 대상을 찾을 수 없습니다.')
            return redirect(f'/warehouse/{warehouse_name}/electric')

        site_name = item[1] or '미입력'

        if request.method == 'POST':
            if not manage_mode:
                conn.close()
                flash('냉방기 추가/삭제는 수정 모드에서만 가능합니다.')
                return redirect(url_for('inspection_units', warehouse_name=warehouse_name, item_id=item_id))

            action = request.form.get('action', '').strip()
            if action == 'add_equipment':
                raw_equipment_no = request.form.get('equipment_no', '').strip().upper()
                parsed = parse_equipment_no(raw_equipment_no)
                if not parsed:
                    conn.close()
                    flash('설비번호 형식이 올바르지 않습니다. 예: A05-01')
                    return redirect(url_for('inspection_units', warehouse_name=warehouse_name, item_id=item_id, mode='edit'))

                equipment_no = format_equipment_no(parsed[1], parsed[0])
                cursor.execute(
                    '''INSERT INTO inspection_units (inventory_id, equipment_no, created_by)
                       VALUES (%s, %s, %s)
                       ON CONFLICT (inventory_id, equipment_no) DO NOTHING''',
                    (item_id, equipment_no, session.get('user_name', '미상'))
                )
                conn.commit()
                conn.close()
                flash(f'설비번호 {equipment_no}가 추가되었습니다.')
                return redirect(url_for('inspection_units', warehouse_name=warehouse_name, item_id=item_id, mode='edit'))

            if action == 'delete_equipment':
                raw_equipment_no = request.form.get('equipment_no', '').strip().upper()
                parsed = parse_equipment_no(raw_equipment_no)
                if not parsed:
                    conn.close()
                    flash('삭제할 설비번호 형식이 올바르지 않습니다.')
                    return redirect(url_for('inspection_units', warehouse_name=warehouse_name, item_id=item_id, mode='edit'))

                equipment_no = format_equipment_no(parsed[1], parsed[0])
                cursor.execute(
                    '''DELETE FROM inspection_photos
                       WHERE record_id IN (
                           SELECT id
                           FROM inspection_records
                           WHERE inventory_id = %s
                             AND equipment_no = %s
                       )''',
                    (item_id, equipment_no)
                )
                cursor.execute(
                    '''DELETE FROM inspection_records
                       WHERE inventory_id = %s
                         AND equipment_no = %s''',
                    (item_id, equipment_no)
                )
                cursor.execute(
                    '''DELETE FROM inspection_units
                       WHERE inventory_id = %s
                         AND equipment_no = %s''',
                    (item_id, equipment_no)
                )
                conn.commit()
                conn.close()
                flash(f'설비번호 {equipment_no}가 삭제되었습니다.')
                return redirect(url_for('inspection_units', warehouse_name=warehouse_name, item_id=item_id, mode='edit'))

        cursor.execute(
            '''SELECT equipment_no
               FROM inspection_units
               WHERE inventory_id = %s''',
            (item_id,)
        )
        equipment_numbers = sort_equipment_nos([row[0] for row in cursor.fetchall() if row[0]])
        suggested_equipment_no = get_next_equipment_no(equipment_numbers)
        suggested_info = parse_equipment_no(suggested_equipment_no)
        suggested_index = suggested_info[1] if suggested_info else DEFAULT_EQUIPMENT_START

        status_map = {}
        cursor.execute(
            '''SELECT u.equipment_no, r.inspector_name, r.inspected_at
               FROM inspection_units u
               LEFT JOIN (
                   SELECT DISTINCT ON (equipment_no)
                          equipment_no,
                          inspector_name,
                          inspected_at
                   FROM inspection_records
                   WHERE inventory_id = %s
                     AND equipment_no IS NOT NULL
                     AND TRIM(equipment_no) <> ''
                   ORDER BY equipment_no, inspected_at DESC
               ) r ON r.equipment_no = u.equipment_no
               WHERE u.inventory_id = %s''',
            (item_id, item_id)
        )
        for equipment_no, inspector_name, inspected_at in cursor.fetchall():
            status_map[equipment_no] = (inspector_name or '', inspected_at)

        equipment_rows = []
        for equipment_no in equipment_numbers:
            inspector_name, inspected_at = status_map.get(equipment_no, ('', None))
            if inspected_at and not isinstance(inspected_at, str):
                inspected_at = inspected_at.strftime('%Y-%m-%d %H:%M:%S')
            equipment_rows.append({
                'equipment_no': equipment_no,
                'inspector_name': inspector_name,
                'inspected_at': inspected_at or '',
                'is_completed': bool(inspected_at)
            })

        conn.close()
        return render_template(
            'inspection_units.html',
            warehouse_name=DIY_ACTIVE_LABEL,
            warehouse_slug=warehouse_name,
            item_id=item_id,
            site_name=site_name,
            equipment_rows=equipment_rows,
            suggested_equipment_no=suggested_equipment_no,
            suggested_equipment_index=suggested_index,
            manage_mode=manage_mode
        )
    except Exception as e:
        try:
            if conn:
                conn.rollback()
                conn.close()
        except Exception:
            pass
        print(f"❌ inspection_units 오류: {type(e).__name__}: {str(e)}")
        flash(f'설비번호 페이지 처리 중 오류가 발생했습니다: {str(e)}')
        return redirect(f'/warehouse/{warehouse_name}/electric')


@app.route('/warehouse/<warehouse_name>/inspection-unit/<int:item_id>/<equipment_no>', methods=['GET', 'POST'])
def inspection_detail_by_equipment(warehouse_name, item_id, equipment_no):
    if 'user_id' not in session:
        return redirect('/')

    db_warehouse_name = get_db_warehouse_from_slug(warehouse_name)
    if not db_warehouse_name:
        return render_template('preparing.html', warehouse_name=DIY_PREPARING_LABEL)

    parsed_equipment = parse_equipment_no(equipment_no)
    if not parsed_equipment:
        flash('설비번호 형식이 올바르지 않습니다.')
        return redirect(url_for('inspection_units', warehouse_name=warehouse_name, item_id=item_id))
    equipment_no = format_equipment_no(parsed_equipment[1], parsed_equipment[0])

    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            '''SELECT id, part_name
               FROM inventory
               WHERE id = %s AND warehouse = %s AND category = %s''',
            (item_id, db_warehouse_name, DIY_CHECKLIST_CATEGORY)
        )
        item = cursor.fetchone()
        if not item:
            conn.close()
            flash('점검 대상을 찾을 수 없습니다.')
            return redirect(f'/warehouse/{warehouse_name}/electric')

        cursor.execute(
            '''SELECT 1
               FROM inspection_units
               WHERE inventory_id = %s AND equipment_no = %s''',
            (item_id, equipment_no)
        )
        if not cursor.fetchone():
            conn.close()
            flash('등록되지 않은 설비번호입니다.')
            return redirect(url_for('inspection_units', warehouse_name=warehouse_name, item_id=item_id))

        cursor.execute(
            '''SELECT id, site_name, inspector_name, inspected_at, checklist_data, memo
               FROM inspection_records
               WHERE inventory_id = %s
                 AND equipment_no = %s
               ORDER BY inspected_at DESC
               LIMIT 1''',
            (item_id, equipment_no)
        )
        latest_record = cursor.fetchone()

        latest_checklist_by_no = {}
        latest_photos = {}
        if latest_record and latest_record[4]:
            try:
                checklist_raw = latest_record[4]
                checklist_data = checklist_raw if isinstance(checklist_raw, list) else json.loads(checklist_raw)
                for row in checklist_data:
                    checkpoint_no = int(row.get('checkpoint_no', 0))
                    latest_checklist_by_no[checkpoint_no] = row.get('result', 'ok')
            except Exception:
                latest_checklist_by_no = {}

            cursor.execute(
                '''SELECT checkpoint_no, checkpoint_name, phase, filename, file_size, supabase_url
                   FROM inspection_photos
                   WHERE record_id = %s''',
                (latest_record[0],)
            )
            for checkpoint_no, checkpoint_name, phase, filename, file_size, supabase_url in cursor.fetchall():
                if checkpoint_no not in latest_photos:
                    latest_photos[checkpoint_no] = {}
                latest_photos[checkpoint_no][phase] = {
                    'checkpoint_name': checkpoint_name,
                    'filename': filename,
                    'file_size': file_size,
                    'supabase_url': supabase_url
                }

        if request.method == 'POST':
            korea_time = get_korea_time().strftime('%Y-%m-%d %H:%M:%S')
            inspector_name = session.get('user_name', '미상')
            requested_site_name = request.form.get('site_name', '').strip()
            if session.get('is_admin'):
                site_name = requested_site_name or (item[1] or '미입력')
            else:
                site_name = (latest_record[1] if latest_record and latest_record[1] else item[1] or '미입력')
            memo = request.form.get('memo', '').strip()
            edit_mode = request.form.get('edit_mode') == '1'
            record_id = request.form.get('record_id', '').strip()
            update_existing = (
                edit_mode and
                record_id.isdigit() and
                latest_record and
                int(record_id) == latest_record[0]
            )

            checklist_data = []
            photo_rows = []

            for checkpoint_no, checkpoint_name in INSPECTION_ITEMS:
                result_value = request.form.get(f'result_{checkpoint_no}', 'ok')
                if result_value not in {'ok', 'need', 'na'}:
                    result_value = 'ok'
                before_file = request.files.get(f'before_{checkpoint_no}')
                after_file = request.files.get(f'after_{checkpoint_no}')
                existing_before = latest_photos.get(checkpoint_no, {}).get('before') if update_existing else None
                existing_after = latest_photos.get(checkpoint_no, {}).get('after') if update_existing else None

                before_name = None
                before_size = None
                before_url = None
                if before_file and before_file.filename:
                    before_name, before_size, before_url = save_inspection_photo(before_file, item_id, checkpoint_no, 'before')
                elif existing_before:
                    before_name = existing_before['filename']
                    before_size = existing_before['file_size']
                    before_url = existing_before['supabase_url']

                after_name = None
                after_size = None
                after_url = None
                if after_file and after_file.filename:
                    after_name, after_size, after_url = save_inspection_photo(after_file, item_id, checkpoint_no, 'after')
                elif existing_after:
                    after_name = existing_after['filename']
                    after_size = existing_after['file_size']
                    after_url = existing_after['supabase_url']

                checklist_data.append({
                    'checkpoint_no': checkpoint_no,
                    'checkpoint_name': checkpoint_name,
                    'result': result_value
                })

                if before_url:
                    photo_rows.append((checkpoint_no, checkpoint_name, 'before', before_name, before_size, before_url))
                if after_url:
                    photo_rows.append((checkpoint_no, checkpoint_name, 'after', after_name, after_size, after_url))

            if update_existing:
                record_id_int = int(record_id)
                cursor.execute(
                    '''UPDATE inspection_records
                       SET site_name = %s,
                           equipment_no = %s,
                           inspector_name = %s,
                           inspected_at = %s,
                           status = %s,
                           checklist_data = %s,
                           memo = %s
                       WHERE id = %s AND inventory_id = %s''',
                    (site_name, equipment_no, inspector_name, korea_time, '?묒뾽 ?꾨즺',
                     json.dumps(checklist_data, ensure_ascii=False), memo, record_id_int, item_id)
                )
                cursor.execute('DELETE FROM inspection_photos WHERE record_id = %s', (record_id_int,))
                target_record_id = record_id_int
            else:
                cursor.execute(
                    '''INSERT INTO inspection_records
                       (inventory_id, warehouse, site_name, equipment_no, inspector_name, inspected_at, status, checklist_data, memo)
                       VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                       RETURNING id''',
                    (item_id, db_warehouse_name, site_name, equipment_no, inspector_name, korea_time, '?묒뾽 ?꾨즺',
                     json.dumps(checklist_data, ensure_ascii=False), memo)
                )
                target_record_id = cursor.fetchone()[0]

            for checkpoint_no, checkpoint_name, phase, filename, file_size, supabase_url in photo_rows:
                cursor.execute(
                    '''INSERT INTO inspection_photos
                       (record_id, checkpoint_no, checkpoint_name, phase, filename, file_size, supabase_url)
                       VALUES (%s, %s, %s, %s, %s, %s, %s)''',
                    (target_record_id, checkpoint_no, checkpoint_name, phase, filename, file_size, supabase_url)
                )

            cursor.execute(
                'UPDATE inventory SET part_name = %s, last_modifier = %s, last_modified = %s WHERE id = %s',
                (site_name, inspector_name, korea_time, item_id)
            )
            cursor.execute(
                '''INSERT INTO inventory_history
                   (inventory_id, change_type, quantity_change, modifier_name, modified_at)
                   VALUES (%s, %s, %s, %s, %s)''',
                (item_id, 'inspection', 0, inspector_name, korea_time)
            )

            conn.commit()
            conn.close()
            flash('점검 내용이 저장되었습니다.')
            return redirect(url_for('inspection_units', warehouse_name=warehouse_name, item_id=item_id))

        editable = (request.args.get('mode') == 'edit') or (latest_record is None)
        latest_record_dict = None
        inspected_at_str = ''
        if latest_record:
            inspected_at_value = latest_record[3]
            if inspected_at_value and not isinstance(inspected_at_value, str):
                inspected_at_str = inspected_at_value.strftime('%Y-%m-%d %H:%M:%S')
            else:
                inspected_at_str = inspected_at_value or ''

            latest_record_dict = {
                'id': latest_record[0],
                'site_name': latest_record[1] or item[1] or '미입력',
                'inspector_name': latest_record[2] or '',
                'inspected_at': inspected_at_str,
                'memo': latest_record[5] or ''
            }

        conn.close()
        return render_template(
            'inspection_detail.html',
            warehouse_name=get_display_warehouse_from_slug(warehouse_name),
            warehouse_slug=warehouse_name,
            item_id=item[0],
            site_name=(latest_record_dict['site_name'] if latest_record_dict else (item[1] or '미입력')),
            equipment_no=equipment_no,
            equipment_selector_url=url_for('inspection_units', warehouse_name=warehouse_name, item_id=item_id),
            back_url=url_for('inspection_units', warehouse_name=warehouse_name, item_id=item_id),
            back_button_label='설비 목록으로',
            show_management_actions=True,
            edit_url=url_for('inspection_detail_by_equipment', warehouse_name=warehouse_name, item_id=item_id, equipment_no=equipment_no, mode='edit'),
            inspector_name=session.get('user_name', '미상'),
            inspection_items=INSPECTION_ITEMS,
            editable=editable,
            latest_record=latest_record_dict,
            latest_checklist_by_no=latest_checklist_by_no,
            latest_photos=latest_photos,
            is_admin=session.get('is_admin', False)
        )
    except Exception as e:
        try:
            if conn:
                conn.rollback()
                conn.close()
        except Exception:
            pass
        print(f"❌ inspection_detail_by_equipment 오류: {type(e).__name__}: {str(e)}")
        flash(f'점검 페이지 처리 중 오류가 발생했습니다: {str(e)}')
        return redirect(url_for('inspection_units', warehouse_name=warehouse_name, item_id=item_id))


@app.route('/warehouse/<warehouse_name>/inspection-export-v2')
def export_inspection_report_v2(warehouse_name):
    if 'user_id' not in session:
        return redirect('/')

    if not OPENPYXL_AVAILABLE:
        flash('엑셀 내보내기 라이브러리가 설치되지 않았습니다. 서버 관리자에게 문의해주세요.')
        return redirect(f'/warehouse/{warehouse_name}/electric')

    db_warehouse_name = get_db_warehouse_from_slug(warehouse_name)
    if not db_warehouse_name:
        return render_template('preparing.html', warehouse_name=DIY_PREPARING_LABEL)

    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            '''SELECT id, part_name
               FROM inventory
               WHERE warehouse = %s AND category = %s
               ORDER BY id''',
            (db_warehouse_name, DIY_CHECKLIST_CATEGORY)
        )
        targets = cursor.fetchall()

        report_rows = []
        for item_id, fallback_site_name in targets:
            cursor.execute(
                '''SELECT equipment_no
                   FROM inspection_units
                   WHERE inventory_id = %s''',
                (item_id,)
            )
            equipment_numbers = [row[0] for row in cursor.fetchall() if row[0]]
            for equipment_no in sort_equipment_nos(equipment_numbers):
                cursor.execute(
                    '''SELECT id, site_name, checklist_data
                       FROM inspection_records
                       WHERE inventory_id = %s
                         AND equipment_no = %s
                       ORDER BY inspected_at DESC
                       LIMIT 1''',
                    (item_id, equipment_no)
                )
                latest_record = cursor.fetchone()
                checklist_by_no = {}
                photos_by_no = defaultdict(dict)
                site_display_name = fallback_site_name or '미입력'

                if latest_record:
                    record_id = latest_record[0]
                    raw_checklist = latest_record[2]
                    if raw_checklist:
                        try:
                            checklist_rows = raw_checklist if isinstance(raw_checklist, list) else json.loads(raw_checklist)
                            for row in checklist_rows:
                                checklist_by_no[int(row.get('checkpoint_no', 0))] = row.get('result', '')
                        except Exception:
                            checklist_by_no = {}

                    cursor.execute(
                        '''SELECT checkpoint_no, phase, supabase_url
                           FROM inspection_photos
                           WHERE record_id = %s''',
                        (record_id,)
                    )
                    for checkpoint_no, phase, photo_url in cursor.fetchall():
                        photos_by_no[int(checkpoint_no)][phase] = photo_url

                report_rows.append({
                    'site_name': site_display_name,
                    'equipment_no': equipment_no,
                    'checklist': checklist_by_no,
                    'photos': photos_by_no
                })

        conn.close()

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = safe_excel_sheet_title('DIY_Report')

        header_fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
        header_font = Font(bold=True)
        thin_side = Side(border_style='thin', color='D0D0D0')
        thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)

        item_titles = {
            1: '고무패킹 교체',
            2: '실내기 Reset',
            3: 'V벨트 교체',
            4: '릴레이 점검',
            5: '배수관 청소',
            6: 'RMS 연동확인',
            7: '자연공조 점검',
            8: '정전 보상',
            9: '실외기 팬/소음',
            10: '송풍기 풍량',
            11: '열화상 측정'
        }

        sheet.cell(row=1, column=1, value='순번')
        sheet.cell(row=1, column=2, value='국사명')
        sheet.cell(row=1, column=3, value='설비번호')
        for checkpoint_no, _checkpoint_name in INSPECTION_ITEMS:
            col_no = 3 + checkpoint_no
            sheet.cell(row=1, column=col_no, value=item_titles.get(checkpoint_no, f'{checkpoint_no}번'))

        max_col = 3 + len(INSPECTION_ITEMS)
        for col_no in range(1, max_col + 1):
            cell = sheet.cell(row=1, column=col_no)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        sheet.column_dimensions['A'].width = 8
        sheet.column_dimensions['B'].width = 34
        sheet.column_dimensions['C'].width = 14
        for col_no in range(4, max_col + 1):
            sheet.column_dimensions[get_column_letter(col_no)].width = 34

        image_refs = []
        data_row_start = 2
        for idx, row_data in enumerate(report_rows, start=1):
            block_start = data_row_start + (idx - 1) * 3
            method_row = block_start
            before_row = block_start + 1
            after_row = block_start + 2

            checklist_map = row_data.get('checklist', {})
            photos_map = row_data.get('photos', {})

            sheet.merge_cells(start_row=method_row, start_column=1, end_row=after_row, end_column=1)
            sheet.merge_cells(start_row=method_row, start_column=2, end_row=after_row, end_column=2)
            sheet.merge_cells(start_row=method_row, start_column=3, end_row=after_row, end_column=3)
            sheet.cell(row=method_row, column=1, value=idx)
            sheet.cell(row=method_row, column=2, value=row_data.get('site_name', '미입력'))
            sheet.cell(row=method_row, column=3, value=row_data.get('equipment_no', '-'))

            sheet.row_dimensions[method_row].height = 95
            sheet.row_dimensions[before_row].height = 120
            sheet.row_dimensions[after_row].height = 120

            for checkpoint_no, _checkpoint_name in INSPECTION_ITEMS:
                col_no = 3 + checkpoint_no
                result_text = checklist_map.get(checkpoint_no, '')
                if result_text == 'ok':
                    result_text = '정상'
                elif result_text == 'need':
                    result_text = '조치필요'
                elif result_text == 'na':
                    result_text = '대상아님'

                method_cell = sheet.cell(row=method_row, column=col_no)
                method_cell.value = f"{INSPECTION_METHOD_GUIDE.get(checkpoint_no, '-')}\n\n결과: {result_text or '-'}"
                method_cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

                before_url = photos_map.get(checkpoint_no, {}).get('before')
                after_url = photos_map.get(checkpoint_no, {}).get('after')

                if not add_excel_image(sheet, before_row, col_no, fetch_image_bytes(before_url), image_refs):
                    sheet.cell(row=before_row, column=col_no, value='작업전 미등록')
                if not add_excel_image(sheet, after_row, col_no, fetch_image_bytes(after_url), image_refs):
                    sheet.cell(row=after_row, column=col_no, value='작업후 미등록')

            for row_no in (method_row, before_row, after_row):
                for col_no in range(1, max_col + 1):
                    sheet.cell(row=row_no, column=col_no).border = thin_border

        sheet.freeze_panes = 'D2'

        output = io.BytesIO()
        workbook.save(output)
        output.seek(0)
        filename = f"SK오앤에스_DIY점검리포트_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        encoded_filename = urllib.parse.quote(filename, safe='')
        return Response(
            output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
    except Exception as e:
        try:
            if conn:
                conn.close()
        except Exception:
            pass
        print(f"❌ export_inspection_report_v2 오류: {type(e).__name__}: {str(e)}")
        flash(f'점검 엑셀 내보내기 중 오류가 발생했습니다: {str(e)}')
        return redirect(f'/warehouse/{warehouse_name}/electric')


@app.route('/warehouse/<warehouse_name>/inspection-export')
def export_inspection_report(warehouse_name):
    """DIY 점검 결과 엑셀 내보내기 (이미지 샘플 구조)"""
    if 'user_id' not in session:
        return redirect('/')

    if not OPENPYXL_AVAILABLE:
        flash('엑셀 내보내기 라이브러리가 설치되지 않았습니다. 서버 관리자에게 문의해주세요.')
        return redirect(f'/warehouse/{warehouse_name}/electric')

    db_warehouse_name = get_db_warehouse_from_slug(warehouse_name)
    if not db_warehouse_name:
        return render_template('preparing.html', warehouse_name=DIY_PREPARING_LABEL)

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            '''SELECT id, part_name
               FROM inventory
               WHERE warehouse = %s AND category = %s
               ORDER BY id''',
            (db_warehouse_name, DIY_CHECKLIST_CATEGORY)
        )
        targets = cursor.fetchall()

        record_map = {}
        for item_id, site_name in targets:
            cursor.execute(
                '''SELECT id, site_name, checklist_data
                   FROM inspection_records
                   WHERE inventory_id = %s
                   ORDER BY inspected_at DESC
                   LIMIT 1''',
                (item_id,)
            )
            latest_record = cursor.fetchone()
            checklist_by_no = {}
            photos_by_no = defaultdict(dict)

            if latest_record:
                record_id = latest_record[0]
                raw_checklist = latest_record[2]
                if raw_checklist:
                    try:
                        checklist_rows = raw_checklist if isinstance(raw_checklist, list) else json.loads(raw_checklist)
                        for row in checklist_rows:
                            checklist_by_no[int(row.get('checkpoint_no', 0))] = row.get('result', '')
                    except Exception:
                        checklist_by_no = {}

                cursor.execute(
                    '''SELECT checkpoint_no, phase, supabase_url
                       FROM inspection_photos
                       WHERE record_id = %s''',
                    (record_id,)
                )
                for checkpoint_no, phase, photo_url in cursor.fetchall():
                    photos_by_no[int(checkpoint_no)][phase] = photo_url

                site_display_name = latest_record[1] or site_name or '미입력'
            else:
                site_display_name = site_name or '미입력'

            record_map[item_id] = {
                'site_name': site_display_name,
                'checklist': checklist_by_no,
                'photos': photos_by_no
            }

        conn.close()

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = safe_excel_sheet_title('DIY_Report')

        header_fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
        header_font = Font(bold=True)
        thin_side = Side(border_style='thin', color='D0D0D0')
        thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)

        item_titles = {
            1: '①고무패킹교체',
            2: '②실내기 Reset',
            3: '③V벨트 교체',
            4: '④타이머릴레이',
            5: '⑤배수관청소',
            6: '⑥RMS 온도센싱',
            7: '⑦자연공조점검',
            8: '⑧정전보상',
            9: '⑨실외기 핀,넝쿨',
            10: '⑩송풍구 풍량',
            11: '⑪열화상 측정'
        }

        sheet.cell(row=1, column=1, value='순번')
        sheet.cell(row=1, column=2, value='국사명')
        for checkpoint_no, _checkpoint_name in INSPECTION_ITEMS:
            col_no = 2 + checkpoint_no
            sheet.cell(row=1, column=col_no, value=item_titles.get(checkpoint_no, f'{checkpoint_no}번'))

        max_col = 2 + len(INSPECTION_ITEMS)
        for col_no in range(1, max_col + 1):
            cell = sheet.cell(row=1, column=col_no)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        sheet.column_dimensions['A'].width = 8
        sheet.column_dimensions['B'].width = 38
        for col_no in range(3, max_col + 1):
            sheet.column_dimensions[get_column_letter(col_no)].width = 34

        image_refs = []
        data_row_start = 2

        for idx, (item_id, _site_name) in enumerate(targets, start=1):
            block_start = data_row_start + (idx - 1) * 3
            method_row = block_start
            before_row = block_start + 1
            after_row = block_start + 2

            row_data = record_map.get(item_id, {})
            checklist_map = row_data.get('checklist', {})
            photos_map = row_data.get('photos', {})

            sheet.merge_cells(start_row=method_row, start_column=1, end_row=after_row, end_column=1)
            sheet.merge_cells(start_row=method_row, start_column=2, end_row=after_row, end_column=2)
            sheet.cell(row=method_row, column=1, value=idx)
            sheet.cell(row=method_row, column=2, value=row_data.get('site_name', '미입력'))
            sheet.cell(row=method_row, column=1).alignment = Alignment(horizontal='center', vertical='center')
            sheet.cell(row=method_row, column=2).alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)

            sheet.row_dimensions[method_row].height = 95
            sheet.row_dimensions[before_row].height = 120
            sheet.row_dimensions[after_row].height = 120

            for checkpoint_no, _checkpoint_name in INSPECTION_ITEMS:
                col_no = 2 + checkpoint_no
                method_text = INSPECTION_METHOD_GUIDE.get(checkpoint_no, '-')
                result_text = checklist_map.get(checkpoint_no, '')
                if result_text == 'ok':
                    result_text = '정상'
                elif result_text == 'need':
                    result_text = '조치필요'
                elif result_text == 'na':
                    result_text = '대상아님'

                method_cell = sheet.cell(row=method_row, column=col_no)
                method_cell.value = f"{method_text}\n\n결과: {result_text or '-'}"
                method_cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

                before_url = photos_map.get(checkpoint_no, {}).get('before')
                after_url = photos_map.get(checkpoint_no, {}).get('after')

                if not add_excel_image(sheet, before_row, col_no, fetch_image_bytes(before_url), image_refs):
                    sheet.cell(row=before_row, column=col_no, value='작업전 미등록').alignment = Alignment(horizontal='center', vertical='center')
                if not add_excel_image(sheet, after_row, col_no, fetch_image_bytes(after_url), image_refs):
                    sheet.cell(row=after_row, column=col_no, value='작업후 미등록').alignment = Alignment(horizontal='center', vertical='center')

            for row_no in (method_row, before_row, after_row):
                for col_no in range(1, max_col + 1):
                    sheet.cell(row=row_no, column=col_no).border = thin_border

        sheet.freeze_panes = 'C2'

        output = io.BytesIO()
        workbook.save(output)
        output.seek(0)

        filename = f'SK오앤에스_DIY점검리포트_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        encoded_filename = urllib.parse.quote(filename, safe="")
        return Response(
            output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        print(f"❌ export_inspection_report 오류: {type(e).__name__}: {str(e)}")
        flash(f'점검 엑셀 내보내기 중 오류가 발생했습니다: {str(e)}')
        return redirect(f'/warehouse/{warehouse_name}/electric')


@app.route('/inspection-method')
def inspection_method():
    if 'user_id' not in session:
        return redirect('/')

    latest_image = get_latest_inspection_method_image()
    image_url = latest_image['public_url'] if latest_image else None

    # 하위 호환: DB가 비어있는 경우 로컬 파일 표시
    if not image_url:
        local_relpath = find_local_inspection_method_image_relpath()
        if local_relpath:
            image_url = url_for('static', filename=local_relpath)

    return render_template(
        'inspection_method.html',
        has_image=bool(image_url),
        image_url=image_url,
        is_admin=session.get('is_admin', False),
        image_meta=latest_image
    )

@app.route('/inspection-method/upload', methods=['POST'])
def upload_inspection_method():
    if 'user_id' not in session:
        return redirect('/')
    if not session.get('is_admin'):
        flash('관리자만 점검 방법 이미지를 등록할 수 있습니다.')
        return redirect(url_for('inspection_method'))

    file_obj = request.files.get('inspection_method_image')
    if not file_obj or file_obj.filename == '':
        flash('업로드할 이미지 파일을 선택해주세요.')
        return redirect(url_for('inspection_method'))
    if not allowed_file(file_obj.filename):
        flash('png/jpg/jpeg/webp 형식만 등록할 수 있습니다.')
        return redirect(url_for('inspection_method'))

    conn = None
    try:
        file_obj.seek(0)
        compressed_bytes, _final_size_kb = compress_image_to_target_size(
            file_obj,
            max_size_mb=2.9,
            max_width=2800,
            quality=92
        )
        if not compressed_bytes:
            flash('점검 방법 이미지 압축에 실패했습니다.')
            return redirect(url_for('inspection_method'))

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        file_key = f"{INSPECTION_METHOD_PREFIX}/inspection_method_{timestamp}_{uuid.uuid4().hex}.jpg"
        public_url = upload_to_supabase_storage(
            compressed_bytes,
            file_key,
            bucket_name=INSPECTION_METHOD_BUCKET,
            content_type='image/jpeg'
        )
        if not public_url:
            flash('점검 방법 이미지 업로드에 실패했습니다.')
            return redirect(url_for('inspection_method'))

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('UPDATE inspection_method_images SET is_active = 0 WHERE is_active = 1')
        cursor.execute(
            '''INSERT INTO inspection_method_images
               (storage_path, public_url, uploaded_by, is_active)
               VALUES (%s, %s, %s, %s)''',
            (file_key, public_url, session.get('user_name', '관리자'), 1)
        )
        conn.commit()
        conn.close()

        flash('점검 방법 이미지가 Supabase에 등록되었습니다.')
        return redirect(url_for('inspection_method'))
    except Exception as e:
        try:
            if conn:
                conn.rollback()
                conn.close()
        except Exception:
            pass
        print(f"❌ upload_inspection_method 오류: {type(e).__name__}: {str(e)}")
        flash(f'점검 방법 이미지 등록 중 오류가 발생했습니다: {str(e)}')
        return redirect(url_for('inspection_method'))


def facility_result_label(value):
    return {'fit': '적합', 'unfit': '부적합', 'na': '해당없음'}.get(value, '-')


def build_facility_section_detail_text(section, state):
    fields = state.get('fields', {}) if isinstance(state, dict) else {}
    parts = []
    for field in section.get('detail_fields', []):
        label = field.get('label', '').strip()
        value = str(fields.get(field.get('key'), '') or '').strip()
        if label:
            parts.append(f"{label}: {value if value else '-'}")
    note = str(state.get('detail_note', '') or '').strip() if isinstance(state, dict) else ''
    if note:
        parts.append(f"비고: {note}")
    return " / ".join(parts) if parts else "-"


def add_picture_to_docx_cell(cell, image_bytes, width_inches=2.7):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if not image_bytes:
        paragraph.add_run('사진 없음')
        return
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=Inches(width_inches))


@app.route('/warehouse/<warehouse_name>/inspection-hangul/<int:item_id>')
def export_facility_report_docx(warehouse_name, item_id):
    if 'user_id' not in session:
        return redirect('/')

    if warehouse_name != DIY_FACILITY_SLUG:
        flash('시설물 적합조사 메뉴에서만 한글 다운로드가 가능합니다.')
        return redirect(url_for('electric_inventory', warehouse_name=warehouse_name))

    if not PYDOCX_AVAILABLE:
        flash('한글 문서 생성을 위한 라이브러리가 설치되지 않았습니다. 관리자에게 문의해주세요.')
        return redirect(url_for('electric_inventory', warehouse_name=warehouse_name))

    conn = None
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
            '''SELECT id, part_name
               FROM inventory
               WHERE id = %s AND warehouse = %s AND category = %s''',
            (item_id, DB_FACILITY_WAREHOUSE, DIY_CHECKLIST_CATEGORY)
        )
        item = cursor.fetchone()
        if not item:
            conn.close()
            flash('점검 대상을 찾을 수 없습니다.')
            return redirect(url_for('electric_inventory', warehouse_name=warehouse_name))

        cursor.execute(
            '''SELECT id, site_name, inspector_name, inspected_at, checklist_data, memo
               FROM inspection_records
               WHERE inventory_id = %s
               ORDER BY inspected_at DESC
               LIMIT 1''',
            (item_id,)
        )
        latest_record = cursor.fetchone()
        if not latest_record:
            conn.close()
            flash('저장된 점검 결과가 없습니다.')
            return redirect(url_for('electric_inventory', warehouse_name=warehouse_name))

        payload = parse_facility_payload(latest_record[4])
        capture_photos = fetch_facility_capture_photos(cursor, latest_record[0])
        conn.close()

        doc = Document()
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Malgun Gothic'
        normal_style.font.size = Pt(10)

        title = doc.add_paragraph('방송통신설비 기술기준 적합 조사ㆍ시험 평가표')
        title.runs[0].bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        inspected_date = payload.get('inspection_date') or (
            latest_record[3].strftime('%Y-%m-%d')
            if latest_record[3] and not isinstance(latest_record[3], str)
            else (latest_record[3] or '-')
        )

        info_table = doc.add_table(rows=1, cols=1)
        info_table.style = 'Table Grid'
        info_table.cell(0, 0).text = (
            f"□ 사업자명: {payload.get('business_name') or '-'}    "
            f"□ 국사명: {latest_record[1] or item[1] or '미입력'}    "
            f"□ 국사형태: {payload.get('site_type') or '-'}    "
            f"□ 점검일자: {inspected_date}"
        )

        doc.add_paragraph('')
        checklist_table = doc.add_table(rows=1, cols=3)
        checklist_table.style = 'Table Grid'
        checklist_table.cell(0, 0).text = '조 사 항 목'
        checklist_table.cell(0, 1).text = '결 과'
        checklist_table.cell(0, 2).text = '비 고'

        for section in FACILITY_SURVEY_SECTIONS:
            code = section['code']
            state = payload['sections'].get(code, {})
            checks = state.get('checks', [])

            section_row = checklist_table.add_row().cells
            section_row[0].text = f"{code}. {section['title']}"
            section_row[1].text = facility_result_label(state.get('result', ''))
            section_row[2].text = ''

            for idx, check_name in enumerate(section.get('checks', [])):
                check_row = checklist_table.add_row().cells
                check_row[0].text = f"- {check_name}"
                check_row[1].text = facility_result_label(checks[idx] if idx < len(checks) else '')
                check_row[2].text = ''

            detail_row = checklist_table.add_row().cells
            detail_row[0].text = f"o 점검결과(상세): {build_facility_section_detail_text(section, state)}"
            detail_row[1].text = ''
            detail_row[2].text = ''

        photo_title = doc.add_paragraph('서면(자가)점검 사진촬영 자료')
        photo_title.runs[0].bold = True

        photo_table = doc.add_table(rows=0, cols=4)
        photo_table.style = 'Table Grid'
        pair_rows = [
            (FACILITY_CAPTURE_SLOTS[0], FACILITY_CAPTURE_SLOTS[1]),
            (FACILITY_CAPTURE_SLOTS[2], FACILITY_CAPTURE_SLOTS[3]),
            (FACILITY_CAPTURE_SLOTS[4], FACILITY_CAPTURE_SLOTS[5]),
            (FACILITY_CAPTURE_SLOTS[6], FACILITY_CAPTURE_SLOTS[7]),
        ]

        for left_slot, right_slot in pair_rows:
            row_title = photo_table.add_row().cells
            row_title[0].text = '점검내용'
            row_title[1].text = f"{left_slot['slot_no']}. {left_slot['title']}"
            row_title[2].text = '점검내용'
            row_title[3].text = f"{right_slot['slot_no']}. {right_slot['title']}"

            row_detail = photo_table.add_row().cells
            row_detail[0].text = '세부사항'
            row_detail[1].text = left_slot['detail']
            row_detail[2].text = '세부사항'
            row_detail[3].text = right_slot['detail']

            row_photo = photo_table.add_row().cells
            left_image_cell = row_photo[0].merge(row_photo[1])
            right_image_cell = row_photo[2].merge(row_photo[3])

            left_photo = capture_photos.get(left_slot['slot_no'])
            right_photo = capture_photos.get(right_slot['slot_no'])
            left_bytes = fetch_image_bytes(left_photo['supabase_url']) if left_photo and left_photo.get('supabase_url') else None
            right_bytes = fetch_image_bytes(right_photo['supabase_url']) if right_photo and right_photo.get('supabase_url') else None
            add_picture_to_docx_cell(left_image_cell, left_bytes, width_inches=2.8)
            add_picture_to_docx_cell(right_image_cell, right_bytes, width_inches=2.8)

        memo_text = latest_record[5] or ''
        doc.add_paragraph(f"메모: {memo_text if memo_text else '-'}")
        doc.add_paragraph(f"점검자: {latest_record[2] or '-'}")

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        site_name_for_file = (latest_record[1] or item[1] or '시설물').replace('/', '_').replace('\\', '_')
        filename = f"서면자가점검_{site_name_for_file}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        encoded_filename = urllib.parse.quote(filename, safe='')

        return Response(
            output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
    except Exception as e:
        try:
            if conn:
                conn.close()
        except Exception:
            pass
        print(f"❌ export_facility_report_docx 오류: {type(e).__name__}: {str(e)}")
        flash(f'한글 다운로드 중 오류가 발생했습니다: {str(e)}')
        return redirect(url_for('electric_inventory', warehouse_name=warehouse_name))

@app.route('/admin/sites', methods=['GET', 'POST'])
def admin_sites():
    """관리자용 국사명 관리"""
    if 'user_id' not in session:
        flash('로그인이 필요합니다.')
        return redirect('/')

    if not session.get('is_admin'):
        flash('관리자 권한이 필요합니다.')
        return redirect('/dashboard')

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        if request.method == 'POST':
            action = request.form.get('action', '').strip()
            site_name = request.form.get('site_name', '').strip()
            site_id = request.form.get('site_id', '').strip()
            korea_time = get_korea_time().strftime('%Y-%m-%d %H:%M:%S')
            admin_name = session.get('user_name', '관리자')

            if action == 'add':
                if not site_name:
                    flash('국사명을 입력해주세요.')
                else:
                    cursor.execute(
                        '''SELECT id
                           FROM inventory
                           WHERE warehouse = %s AND category = %s AND part_name = %s''',
                        (DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY, site_name)
                    )
                    if cursor.fetchone():
                        flash('이미 등록된 국사명입니다.')
                    else:
                        cursor.execute(
                            '''INSERT INTO inventory
                               (warehouse, category, part_name, quantity, last_modifier, last_modified)
                               VALUES (%s, %s, %s, %s, %s, %s)''',
                            (DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY, site_name, 0, admin_name, korea_time)
                        )
                        conn.commit()
                        flash('국사명이 추가되었습니다.')

            elif action == 'update':
                if not (site_id.isdigit() and site_name):
                    flash('수정할 국사명을 확인해주세요.')
                else:
                    cursor.execute(
                        '''UPDATE inventory
                           SET part_name = %s, last_modifier = %s, last_modified = %s
                           WHERE id = %s AND warehouse = %s AND category = %s''',
                        (site_name, admin_name, korea_time, int(site_id), DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY)
                    )
                    conn.commit()
                    flash('국사명이 수정되었습니다.')

            elif action == 'delete':
                if not site_id.isdigit():
                    flash('삭제할 대상을 확인해주세요.')
                else:
                    cursor.execute(
                        'SELECT COUNT(*) FROM inspection_records WHERE inventory_id = %s',
                        (int(site_id),)
                    )
                    used_count = cursor.fetchone()[0]
                    if used_count > 0:
                        flash('점검 이력이 있는 국사명은 삭제할 수 없습니다. 이름 수정을 이용해주세요.')
                    else:
                        cursor.execute(
                            '''DELETE FROM inventory
                               WHERE id = %s AND warehouse = %s AND category = %s''',
                            (int(site_id), DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY)
                        )
                        conn.commit()
                        flash('국사명이 삭제되었습니다.')

        cursor.execute(
            '''SELECT id, part_name, last_modifier, last_modified
               FROM inventory
               WHERE warehouse = %s AND category = %s
               ORDER BY id''',
            (DB_ACTIVE_WAREHOUSE, DIY_CHECKLIST_CATEGORY)
        )
        raw_rows = cursor.fetchall()
        conn.close()

        site_rows = []
        for row in raw_rows:
            last_modified = row[3]
            if last_modified and not isinstance(last_modified, str):
                last_modified = last_modified.strftime('%Y-%m-%d %H:%M:%S')
            site_rows.append({
                'id': row[0],
                'site_name': row[1],
                'last_modifier': row[2] or '-',
                'last_modified': last_modified or '-'
            })

        return render_template('admin_sites.html', site_rows=site_rows, warehouse_name=DIY_ACTIVE_LABEL)

    except Exception as e:
        try:
            conn.rollback()
            conn.close()
        except Exception:
            pass
        print(f"❌ admin_sites 오류: {type(e).__name__}: {str(e)}")
        flash(f'국사명 관리 중 오류가 발생했습니다: {str(e)}')
        return redirect('/admin/dashboard')

@app.route('/add_inventory_item', methods=['POST'])
def add_inventory_item():
    """재고 아이템 추가 (관리자 전용)"""
    if 'user_id' not in session or not session.get('is_admin'):
        flash('관리자 권한이 필요합니다.')
        return redirect('/')

    warehouse_name = request.form['warehouse_name']
    category = request.form['category']
    part_name = request.form['part_name']
    quantity = int(request.form['quantity'])
    korea_time = get_korea_time().strftime('%Y-%m-%d %H:%M:%S')

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('INSERT INTO inventory (warehouse, category, part_name, quantity, last_modifier, last_modified) VALUES (%s, %s, %s, %s, %s, %s)',
                      (warehouse_name, category, part_name, quantity, session['user_name'], korea_time))
        
        conn.commit()
        conn.close()
        flash('재고 아이템이 추가되었습니다.')
        
    except Exception as e:
        flash('재고 추가 중 오류가 발생했습니다.')
    
    return redirect(f'/warehouse/{get_slug_from_db_warehouse(warehouse_name)}/electric')

@app.route('/update_quantity', methods=['POST'])
def update_quantity():
    """재고 수량 업데이트"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '로그인이 필요합니다.'}), 401

    try:
        data = request.get_json()
        item_id = data['item_id']
        change_type = data['change_type']
        quantity_change = int(data['quantity'])

        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('SELECT quantity, warehouse FROM inventory WHERE id = %s', (item_id,))
        result = cursor.fetchone()
        if not result:
            conn.close()
            return jsonify({'success': False, 'message': '재고 항목을 찾을 수 없습니다.'})
            
        current_quantity, warehouse = result

        if change_type == 'out':
            quantity_change = -quantity_change
            if current_quantity + quantity_change < 0:
                conn.close()
                return jsonify({'success': False, 'message': '재고가 부족합니다.'})

        new_quantity = current_quantity + quantity_change
        korea_time = get_korea_time().strftime('%Y-%m-%d %H:%M:%S')

        cursor.execute('UPDATE inventory SET quantity = %s, last_modifier = %s, last_modified = %s WHERE id = %s',
                      (new_quantity, session['user_name'], korea_time, item_id))

        cursor.execute('INSERT INTO inventory_history (inventory_id, change_type, quantity_change, modifier_name, modified_at) VALUES (%s, %s, %s, %s, %s)',
                      (item_id, change_type, quantity_change, session['user_name'], korea_time))

        conn.commit()
        conn.close()

        return jsonify({'success': True, 'new_quantity': new_quantity})
        
    except Exception as e:
        return jsonify({'success': False, 'message': '수량 업데이트 중 오류가 발생했습니다.'})

@app.route('/upload_photo/<int:item_id>', methods=['POST'])
def upload_photo(item_id):
    """사진 업로드 - Supabase Storage + 이미지 압축"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '로그인이 필요합니다.'}), 401

    if 'photo' not in request.files:
        return jsonify({'success': False, 'message': '파일이 선택되지 않았습니다.'})

    file = request.files['photo']
    if file.filename == '':
        return jsonify({'success': False, 'message': '파일이 선택되지 않았습니다.'})

    if file and allowed_file(file.filename):
        try:
            # 원본 파일 크기 확인
            file.seek(0, 2)  # 파일 끝으로 이동
            original_size_bytes = file.tell()
            file.seek(0)  # 파일 시작으로 이동
            original_size_mb = original_size_bytes / (1024 * 1024)
            
            print(f"📊 원본 이미지 크기: {original_size_mb:.1f}MB")
            
            # 고유 파일명 생성
            filename = f"{uuid.uuid4().hex}_{int(datetime.now().timestamp())}.jpg"
            
            # 이미지 압축 (1MB 미만으로)
            compressed_bytes, final_size_kb = compress_image_to_target_size(
                file, 
                max_size_mb=0.9,  # 1MB보다 약간 작게
                max_width=800,    # 최대 800px 폭
                quality=85        # 초기 품질
            )
            
            if not compressed_bytes:
                return jsonify({'success': False, 'message': '이미지 압축에 실패했습니다.'})
            
            # Supabase Storage에 업로드
            supabase_url = upload_to_supabase_storage(compressed_bytes, filename)
            
            if supabase_url:
                # 데이터베이스에 정보 저장
                conn = get_db_connection()
                cursor = conn.cursor()
                
                cursor.execute('''INSERT INTO photos 
                                (inventory_id, filename, original_name, file_size, uploaded_by, supabase_url) 
                                VALUES (%s, %s, %s, %s, %s, %s)''',
                              (item_id, filename, file.filename, int(final_size_kb), 
                               session['user_name'], supabase_url))
                
                conn.commit()
                conn.close()
                
                return jsonify({
                    'success': True, 
                    'message': f'사진이 업로드되었습니다. (원본: {original_size_mb:.1f}MB → 압축: {final_size_kb:.0f}KB)',
                    'url': supabase_url,
                    'original_size': f"{original_size_mb:.1f}MB",
                    'compressed_size': f"{final_size_kb:.0f}KB"
                })
            else:
                return jsonify({'success': False, 'message': 'Supabase Storage 업로드에 실패했습니다.'})
                
        except Exception as e:
            print(f"❌ 사진 업로드 전체 오류: {e}")
            return jsonify({'success': False, 'message': f'사진 업로드 중 오류가 발생했습니다: {str(e)}'})

    return jsonify({'success': False, 'message': '지원하지 않는 파일 형식입니다.'})

@app.route('/photos/<int:item_id>')
def view_photos(item_id):
    """사진 보기 페이지 - datetime 오류 완전 해결"""
    if 'user_id' not in session:
        return redirect('/')

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('SELECT id, filename, original_name, file_size, uploaded_by, uploaded_at, supabase_url FROM photos WHERE inventory_id = %s ORDER BY uploaded_at DESC', (item_id,))
        raw_photos = cursor.fetchall()
        
        cursor.execute('SELECT part_name, warehouse, category FROM inventory WHERE id = %s', (item_id,))
        item_info = cursor.fetchone()
        conn.close()

        # 🔧 datetime 객체를 문자열로 변환
        photos = []
        for photo in raw_photos:
            photo_list = list(photo)
            if photo_list[5]:  # uploaded_at가 존재하면
                if isinstance(photo_list[5], str):
                    # 이미 문자열이면 그대로 사용
                    pass
                else:
                    # datetime 객체면 문자열로 변환
                    photo_list[5] = photo_list[5].strftime('%Y-%m-%d %H:%M:%S')
            photos.append(photo_list)

        return render_template('photos.html', 
                             photos=photos, 
                             item_id=item_id, 
                             item_info=item_info,
                             is_admin=session.get('is_admin', False))
        
    except Exception as e:
        print(f"❌ 사진 보기 페이지 오류: {type(e).__name__}: {str(e)}")
        # 🔧 리디렉션 대신 오류 페이지 표시
        return f'''
        <html>
        <head><title>사진 보기 오류</title></head>
        <body style="font-family: Arial, sans-serif; padding: 20px; text-align: center;">
            <h2>🔧 사진을 불러오는 중 문제가 발생했습니다</h2>
            <p>오류: {str(e)}</p>
            <a href="javascript:history.back()">← 뒤로가기</a>
        </body>
        </html>
        '''

@app.route('/delete_photo/<int:photo_id>')
def delete_photo(photo_id):
    """사진 삭제 (관리자 전용)"""
    if 'user_id' not in session or not session.get('is_admin'):
        flash('관리자 권한이 필요합니다.')
        return redirect('/')

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('SELECT filename, inventory_id, supabase_url FROM photos WHERE id = %s', (photo_id,))
        photo_info = cursor.fetchone()
        
        if photo_info:
            filename, inventory_id, supabase_url = photo_info
            
            # Supabase Storage에서 파일 삭제 (선택사항)
            if supabase_url:
                try:
                    delete_url = f"{SUPABASE_URL}/storage/v1/object/warehouse-photos/{filename}"
                    headers = {'Authorization': f'Bearer {SUPABASE_SERVICE_KEY}'}
                    requests.delete(delete_url, headers=headers)
                    print(f"✅ Supabase Storage에서 파일 삭제: {filename}")
                except Exception as storage_error:
                    print(f"⚠️ Supabase Storage 파일 삭제 실패: {storage_error}")
            
            # 로컬 파일 삭제 (호환성)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.exists(file_path):
                os.remove(file_path)
            
            cursor.execute('DELETE FROM photos WHERE id = %s', (photo_id,))
            conn.commit()
            flash('사진이 삭제되었습니다.')
            conn.close()
            return redirect(f'/photos/{inventory_id}')
        else:
            flash('삭제할 사진을 찾을 수 없습니다.')
            conn.close()
        
    except Exception as e:
        flash('사진 삭제 중 오류가 발생했습니다.')
    
    if session.get('is_admin'):
        return redirect('/admin/warehouse')
    else:
        return redirect('/dashboard')

@app.route('/search_inventory')
def search_inventory():
    """DIY 점검 대상 검색 페이지."""
    if 'user_id' not in session:
        return redirect('/')

    query = request.args.get('q', '').strip()
    warehouse = request.args.get('warehouse', '').strip()
    warehouse_db = normalize_warehouse_filter(warehouse)

    print(f"🔍 점검 대상 검색 요청: query='{query}', warehouse='{warehouse}'")

    if not query and not warehouse:
        return render_template(
            'search_results.html',
            results=[],
            query='',
            warehouse='',
            warehouse_display_name='',
            is_admin=session.get('is_admin', False)
        )

    # 현재 검색은 DIY 점검 대상만 제공
    if not warehouse_db:
        warehouse_db = DB_ACTIVE_WAREHOUSE
        warehouse = DIY_ACTIVE_SLUG

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        where_conditions = []
        params = []

        if query:
            where_conditions.append("(i.part_name ILIKE %s OR COALESCE(r.inspector_name, '') ILIKE %s)")
            params.append(f'%{query}%')
            params.append(f'%{query}%')

        if warehouse_db:
            where_conditions.append("i.warehouse = %s")
            params.append(warehouse_db)

        where_conditions.append("i.category = %s")
        params.append(DIY_CHECKLIST_CATEGORY)

        where_clause = " AND ".join(where_conditions)

        query_sql = f'''
            SELECT i.id,
                   i.part_name AS site_name,
                   i.warehouse AS warehouse_name,
                   r.inspector_name,
                   r.inspected_at
            FROM inventory i
            LEFT JOIN (
                SELECT DISTINCT ON (inventory_id)
                       inventory_id,
                       inspector_name,
                       inspected_at
                FROM inspection_records
                ORDER BY inventory_id, inspected_at DESC
            ) r ON i.id = r.inventory_id
            WHERE {where_clause}
            ORDER BY i.part_name
        '''

        cursor.execute(query_sql, params)
        raw_results = cursor.fetchall()
        conn.close()

        results = []
        for item_id, site_name, item_warehouse_name, inspector_name, inspected_at in raw_results:
            inspected_at_str = ''
            if inspected_at:
                if isinstance(inspected_at, str):
                    inspected_at_str = inspected_at
                else:
                    inspected_at_str = inspected_at.strftime('%Y-%m-%d %H:%M:%S')

            is_completed = bool(inspected_at_str)
            warehouse_slug = get_slug_from_db_warehouse(item_warehouse_name)
            if uses_equipment_units(warehouse_slug):
                detail_url = url_for('inspection_units', warehouse_name=warehouse_slug, item_id=item_id)
            else:
                detail_url = url_for('inspection_detail', warehouse_name=warehouse_slug, item_id=item_id)
            results.append({
                'id': item_id,
                'site_name': site_name or '미입력',
                'inspector_name': inspector_name or '',
                'inspected_at': inspected_at_str,
                'is_completed': is_completed,
                'action_label': '보기' if is_completed else '점검하기',
                'detail_url': detail_url
            })

        print(f"✅ 검색 결과: {len(results)}개 항목")

        return render_template(
            'search_results.html',
            results=results,
            query=query,
            warehouse=warehouse or DIY_ACTIVE_SLUG,
            warehouse_display_name=get_display_warehouse_from_slug(warehouse or DIY_ACTIVE_SLUG),
            is_admin=session.get('is_admin', False)
        )

    except Exception as e:
        print(f"❌ 검색 중 오류: {type(e).__name__}: {str(e)}")
        return render_template(
            'search_results.html',
            results=[],
            query=query,
            warehouse=warehouse or DIY_ACTIVE_SLUG,
            warehouse_display_name=get_display_warehouse_from_slug(warehouse or DIY_ACTIVE_SLUG),
            is_admin=session.get('is_admin', False),
            error_message=f'검색 중 오류가 발생했습니다: {str(e)}'
        )

@app.route('/delete_inventory/<int:item_id>')
def delete_inventory(item_id):
    """재고 삭제 (관리자 전용)"""
    if 'user_id' not in session or not session.get('is_admin'):
        flash('관리자 권한이 필요합니다.')
        return redirect('/')

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 관련 사진들 삭제
        cursor.execute('SELECT filename, supabase_url FROM photos WHERE inventory_id = %s', (item_id,))
        photos = cursor.fetchall()
        
        for photo in photos:
            filename, supabase_url = photo
            
            # Supabase Storage에서 파일 삭제
            if supabase_url:
                try:
                    delete_url = f"{SUPABASE_URL}/storage/v1/object/warehouse-photos/{filename}"
                    headers = {'Authorization': f'Bearer {SUPABASE_SERVICE_KEY}'}
                    requests.delete(delete_url, headers=headers)
                except Exception as storage_error:
                    print(f"⚠️ Supabase Storage 파일 삭제 실패: {storage_error}")
            
            # 로컬 파일 삭제 (호환성)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.exists(file_path):
                os.remove(file_path)
        
        cursor.execute('DELETE FROM photos WHERE inventory_id = %s', (item_id,))
        cursor.execute('DELETE FROM inventory_history WHERE inventory_id = %s', (item_id,))
        cursor.execute('SELECT warehouse, category FROM inventory WHERE id = %s', (item_id,))
        item_info = cursor.fetchone()
        cursor.execute('DELETE FROM inventory WHERE id = %s', (item_id,))
        
        conn.commit()
        conn.close()
        
        flash('재고 아이템이 삭제되었습니다.')
        
        if item_info:
            warehouse, category = item_info
            if category in [DIY_CHECKLIST_CATEGORY, '전기차']:
                return redirect(f'/warehouse/{get_slug_from_db_warehouse(warehouse)}/electric')
            else:
                return redirect(f'/warehouse/{get_slug_from_db_warehouse(warehouse)}/access')
        
    except Exception as e:
        flash('재고 삭제 중 오류가 발생했습니다.')
    
    if session.get('is_admin'):
        return redirect('/admin/warehouse')
    else:
        return redirect('/dashboard')


@app.route('/delete_receipt/<int:receipt_id>')
def delete_receipt(receipt_id):
    """인수증 삭제 (관리자 전용)"""
    if 'user_id' not in session or not session.get('is_admin'):
        flash('관리자 권한이 필요합니다.')
        return redirect('/')
    
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 인수증 정보 조회 (창고명 확인용)
        cursor.execute('SELECT items_data FROM delivery_receipts WHERE id = %s', (receipt_id,))
        receipt_info = cursor.fetchone()
        
        if receipt_info:
            # 창고명 추출
            warehouse_name = DB_ACTIVE_WAREHOUSE  # 기본값
            try:
                items_data = receipt_info[0]
                if isinstance(items_data, str):
                    parsed_data = json.loads(items_data)
                    if isinstance(parsed_data, dict) and 'warehouse' in parsed_data:
                        warehouse_name = parsed_data['warehouse']
            except:
                pass
            
            # 인수증 삭제
            cursor.execute('DELETE FROM delivery_receipts WHERE id = %s', (receipt_id,))
            conn.commit()
            flash('인수증이 삭제되었습니다.')
            
            conn.close()
            return redirect(f'/receipt_history/{get_slug_from_db_warehouse(warehouse_name)}')
        else:
            flash('삭제할 인수증을 찾을 수 없습니다.')
            conn.close()
        
    except Exception as e:
        print(f"인수증 삭제 오류: {e}")
        flash('인수증 삭제 중 오류가 발생했습니다.')
    
    return redirect('/admin/dashboard')



@app.route('/logout')
def logout():
    """로그아웃"""
    session.clear()
    flash('로그아웃되었습니다.')
    return redirect('/')

@app.route('/inventory_history/<int:item_id>')
def inventory_history(item_id):
    """재고 이력 페이지 - 무한 리디렉션 및 datetime 오류 해결"""
    if 'user_id' not in session:
        return redirect('/')

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 재고 이력 조회
        cursor.execute('''SELECT change_type, quantity_change, modifier_name, modified_at 
                         FROM inventory_history 
                         WHERE inventory_id = %s 
                         ORDER BY modified_at DESC''', (item_id,))
        raw_history = cursor.fetchall()
        
        # 재고 정보 조회
        cursor.execute('SELECT part_name, warehouse, category, quantity FROM inventory WHERE id = %s', (item_id,))
        item_info = cursor.fetchone()
        
        conn.close()
        
        # 🔧 datetime 객체를 문자열로 변환 (오류 방지)
        history = []
        for record in raw_history:
            record_list = list(record)
            if record_list[3]:  # modified_at이 존재하면
                if isinstance(record_list[3], str):
                    # 이미 문자열이면 그대로 사용
                    pass
                else:
                    # datetime 객체면 문자열로 변환
                    record_list[3] = record_list[3].strftime('%Y-%m-%d %H:%M:%S')
            history.append(record_list)
        
        return render_template('inventory_history.html',
                             history=history,
                             item_info=item_info,
                             item_id=item_id)
        
    except Exception as e:
        print(f"❌ 재고 이력 페이지 오류: {type(e).__name__}: {str(e)}")
        
        # 🔧 리디렉션 대신 오류 페이지 표시 (무한 루프 방지)
        return f'''
        <html>
        <head><title>재고 이력 오류</title></head>
        <body style="font-family: Arial, sans-serif; padding: 20px; text-align: center;">
            <h2>🔧 재고 이력을 불러오는 중 문제가 발생했습니다</h2>
            <p>오류: {str(e)}</p>
            <a href="javascript:history.back()">← 뒤로가기</a>
        </body>
        </html>
        '''

@app.route('/export_inventory')
def export_inventory():
    """재고 데이터 내보내기 - 한글 인코딩 문제 완전 해결"""
    if 'user_id' not in session or not session.get('is_admin'):
        flash('관리자 권한이 필요합니다.')
        return redirect('/')
    
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('''SELECT warehouse, category, part_name, quantity, last_modifier, last_modified 
                         FROM inventory 
                         ORDER BY warehouse, category, part_name''')
        inventory_data = cursor.fetchall()
        conn.close()
        
        # 🔧 한글 인코딩 문제 해결: UTF-8 BOM 추가
        output = io.StringIO()
        
        # UTF-8 BOM 추가 (Excel 한글 인식용)
        output.write('\ufeff')  # BOM 추가
        
        writer = csv.writer(output)
        
        # 헤더 작성
        writer.writerow(['창고', '카테고리', '부품명', '수량', '최종수정자', '최종수정일'])
        
        # 데이터 작성
        for row in inventory_data:
            # datetime 객체 처리
            row_list = list(row)
            if row_list[5] and not isinstance(row_list[5], str):
                row_list[5] = row_list[5].strftime('%Y-%m-%d %H:%M:%S')
            writer.writerow(row_list)
        
        # 파일 다운로드 응답 (UTF-8 BOM 포함)
        filename = f'SK오앤에스_재고목록_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
        encoded_filename = urllib.parse.quote(filename, safe="")
        
        response = Response(
            output.getvalue().encode('utf-8-sig'),  # UTF-8 BOM 인코딩
            mimetype='text/csv',
            headers={
                'Content-Disposition': f'attachment; filename*=UTF-8\'\'{encoded_filename}'
            }
        )
        
        return response
        
    except Exception as e:
        flash('데이터 내보내기 중 오류가 발생했습니다.')
        return redirect('/admin/dashboard')

@app.route('/health')
def health():
    """시스템 상태 확인 API"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT 1')
        conn.close()
        
        return jsonify({
            'status': 'healthy',
            'database': 'postgresql',
            'supabase_connected': True,
            'storage_enabled': bool(SUPABASE_URL and SUPABASE_SERVICE_KEY),
            'email_enabled': bool(SMTP_USERNAME and SMTP_PASSWORD),
            'timestamp': datetime.now().isoformat(),
            'message': 'SK오앤에스 DIY System (Supabase PostgreSQL + Storage + Email) 정상 작동 중'
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'database': 'postgresql',
            'supabase_connected': False,
            'timestamp': datetime.now().isoformat(),
            'message': f'Supabase 연결 오류: {str(e)}'
        }), 500

# ========
# 에러 핸들러
# ========
@app.errorhandler(404)
def page_not_found(error):
    """404 에러 핸들러"""
    return '''
    <html>
    <head><title>404 - 페이지를 찾을 수 없음</title></head>
    <body style="font-family: Arial, sans-serif; text-align: center; padding: 50px;">
        <h1>404 - 페이지를 찾을 수 없습니다</h1>
        <p>요청하신 페이지가 존재하지 않습니다.</p>
        <a href="/" style="color: #007bff; text-decoration: none;">홈으로 돌아가기</a>
    </body>
    </html>
    ''', 404

@app.errorhandler(500)
def internal_error(error):
    """500 에러 핸들러"""
    return '''
    <html>
    <head><title>500 - 서버 오류</title></head>
    <body style="font-family: Arial, sans-serif; text-align: center; padding: 50px;">
        <h1>500 - 서버 내부 오류</h1>
        <p>서버에서 문제가 발생했습니다.</p>
        <p>잠시 후 다시 시도해주세요.</p>
        <a href="/" style="color: #007bff; text-decoration: none;">홈으로 돌아가기</a>
    </body>
    </html>
    ''', 500

@app.errorhandler(403)
def forbidden(error):
    """403 에러 핸들러"""
    return '''
    <html>
    <head><title>403 - 접근 권한 없음</title></head>
    <body style="font-family: Arial, sans-serif; text-align: center; padding: 50px;">
        <h1>403 - 접근 권한이 없습니다</h1>
        <p>이 페이지에 접근할 권한이 없습니다.</p>
        <a href="/" style="color: #007bff; text-decoration: none;">홈으로 돌아가기</a>
    </body>
    </html>
    ''', 403

# ========
# 메인 실행 부분
# ========
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 10000))
    is_render = os.environ.get('RENDER') is not None
    
    print("")
    print("🎯 최종 시스템 정보:")
    print(f"📱 포트: {port}")
    print(f"🗄️ 데이터베이스: PostgreSQL (Supabase)")
    print(f"📁 파일 저장: Supabase Storage + 이미지 압축")
    print(f"📧 이메일: {'설정됨' if SMTP_USERNAME else '미설정'}")
    print(f"🔒 보안: 관리자/사용자 권한 분리")
    print(f"🌐 환경: {'Production (Render)' if is_render else 'Development'}")
    print(f"💾 데이터 보존: 영구 (Supabase)")
    print(f"📸 이미지 압축: 10MB → 1MB 미만 자동 압축")
    print(f"📋 인수증 기능: 전자서명 + 이메일 발송")
    print(f"🏪 창고: {', '.join(WAREHOUSES)}")
    print("=" * 60)
    print("🚀 SK오앤에스 DIY System 시작!")
    print("=" * 60)
    
    try:
        app.run(host='0.0.0.0', port=port, debug=not is_render)
    except Exception as e:
        print(f"❌ 서버 시작 실패: {e}")
        sys.exit(1)















